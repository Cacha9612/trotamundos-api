from fastapi import FastAPI, File, UploadFile, HTTPException,Query, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.encoders import jsonable_encoder
from docx.shared import RGBColor

from sqlalchemy import text
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.sql import text
from database import engine

import pandas as pd  # Importa Pandas
import json
import os
import pdfkit
import base64
import logging
import io

from datetime import timedelta, datetime

from typing import List, Optional, Dict

from PIL import Image

from io import BytesIO

from pydantic import BaseModel

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from modelos import (
    GetCliente, ResponseModel, SaveCliente, Vehiculo, GetOrden, GetVehiculo,
    SaveOrden, DatosLogin, Token, OrdenCompleta, Roles, Estatus, SaveUsuario,
    saveVehiculo, ImageData, Empleado, OrdenService, Checklist, CheckListHistorico,
    Flotillas, ModificarVehiculo, Tecnicos, AsignarOrden, ReporteVentas, VehiculoV2
)
ACCESS_TOKEN_EXPIRE_MINUTES = 480
from negocios import Negocios
from utils import utilsclass

# Configuración del logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()

origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
options = {
    'page-size': 'A4',
    'margin-top': '1cm',
    'margin-right': '1cm',
    'margin-bottom': '1cm',
    'margin-left': '1cm',
}

class DocumentRequestV2(BaseModel):
    id_checklist: int  # Nuevo parámetro para identificar el checklist
    placeholders: Dict[str, str]
    logo_base64: str
    logo_derecho_base64: str

# Función para comprimir imágenes
def compress_image(image_path: str, quality: int = 85) -> bytes:
    img = Image.open(image_path)
    img = img.convert("RGB")  # Convertir a RGB si la imagen está en otro formato
    buffer = BytesIO()
    img.save(buffer, format="JPEG", quality=quality)  # Comprimir la imagen
    return buffer.getvalue()

# Función para validar el tamaño de una imagen Base64
def validate_image_size(base64_image: str, max_size_mb: int = 5) -> bool:
    # Recortar prefijo si está presente
    if base64_image.startswith("data:image/"):
        base64_image = base64_image.split(",")[1]

    # Calcular tamaño del contenido
    image_size = len(base64_image) * 3 / 4
    image_size_mb = image_size / (1024 * 1024)
    if image_size_mb > max_size_mb:
        raise ValueError(f"El tamaño de la imagen excede el límite de {max_size_mb} MB.")
    return True

# Función para obtener imágenes desde la base de datos para un checklist específico
def get_service_one(id_checklist: int, num_placeholders: int) -> List[str]:
    query = text("EXEC [dbo].[sp_get_all_checklist_Evidencias] @IdCheckList = :id_checklist")

    try:
        with engine.connect() as connection:
            result = connection.execute(query, {"id_checklist": id_checklist})
            columns = result.keys()
            rows = result.fetchall()
            roles_df = pd.DataFrame(rows, columns=columns)
    except Exception as e:
        logging.error(f"Error ejecutando el procedimiento almacenado: {e}")
        raise HTTPException(status_code=500, detail=f"Error ejecutando el procedimiento almacenado: {e}")

    if roles_df.empty:
        raise HTTPException(status_code=404, detail="No se encontraron datos para el checklist proporcionado.")

    # Filtrar columnas de imágenes
    image_columns = [col for col in roles_df.columns if isinstance(col, str) and '_foto' in col]

    if not image_columns:
        raise ValueError("El procedimiento almacenado no retornó columnas relacionadas con imágenes.")

    # Extraer imágenes de las columnas
    image_list = []
    for col in image_columns:
        image_list.extend(roles_df[col].dropna().tolist())

    # Filtrar imágenes válidas
    image_list = [img for img in image_list if img.strip() != '']

    # Asegurar que el número de imágenes no exceda los placeholders
    if len(image_list) > num_placeholders:
        image_list = image_list[:num_placeholders]

    image_list_base64 = []
    for img in image_list:
        try:
            validate_image_size(img)  # Validar tamaño de la imagen
            image_list_base64.append(img)  # La imagen ya está en formato Base64
        except Exception as e:
            logging.error(f"Error al procesar la imagen: {e}")
            raise ValueError(f"No se pudo procesar la imagen: {e}")

    return image_list_base64


# Función para generar el documento Word con imágenes y placeholders
def generate_word_documentv2(placeholders: Dict[str, str], images_base64: List[str], logo_base64: str, logo_derecho_base64: str) -> BytesIO:
    doc = Document()

    # Crear el encabezado
    section = doc.sections[-1]
    header = section.header
    header_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    header_table.autofit = True

    # Insertar logo izquierdo
    if logo_base64:
        image_data = base64.b64decode(logo_base64)
        image_stream = BytesIO(image_data)
        left_cell = header_table.cell(0, 0)
        paragraph = left_cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(1.5))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Insertar título centrado
    title_cell = header_table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    set_header_format(title_paragraph, "FORMATO DE EVIDENCIAS FOTOGRÁFICAS")

    # Insertar logo derecho
    if logo_derecho_base64:
        image_data_right = base64.b64decode(logo_derecho_base64)
        image_stream_right = BytesIO(image_data_right)
        right_cell = header_table.cell(0, 2)
        paragraph_right = right_cell.paragraphs[0]
        run_right = paragraph_right.add_run()
        run_right.add_picture(image_stream_right, width=Inches(1.5))
        paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Espacio después del encabezado
    doc.add_paragraph()

    # Crear tabla de datos del vehículo
    table_data = doc.add_table(rows=3, cols=4)
    table_data.style = "Table Grid"

    keys = list(placeholders.keys())
    values = list(placeholders.values())

    # Ajustar keys_and_values en caso de que el número de placeholders sea menor que 6
    keys_and_values = [
        (keys[i], values[i]) if i < len(keys) else ("", "") for i in range(6)
    ]

    for i in range(3):  # Llenar filas y columnas
        for j in range(4):
            index = i * 4 + j
            if index < len(keys_and_values):
                key, value = keys_and_values[index]
                cell = table_data.cell(i, j)
                cell.text = f"{key.upper()}: {value}"
                cell.paragraphs[0].runs[0].font.size = Pt(10)

    # Espacio entre tabla y fotos
    doc.add_paragraph()

    # Crear tabla para imágenes (2 columnas)
    num_images = len(images_base64)
    rows = (num_images + 1) // 2 if num_images > 1 else 1  # Asegurar que haya al menos una fila

    table_images = doc.add_table(rows=rows, cols=2)
    table_images.style = "Table Grid"

    # Verificar que haya imágenes base64 antes de intentar insertarlas
    if not images_base64:
        raise ValueError("No se encontraron imágenes válidas para agregar.")

    for idx, image_base64 in enumerate(images_base64):
        image_data = base64.b64decode(image_base64)
        image_stream = BytesIO(image_data)

        row = idx // 2
        col = idx % 2
        cell = table_images.cell(row, col)

        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(2.5), height=Inches(2.5))

    # Guardar documento en BytesIO
    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)

    return word_stream

@app.post("/generate_and_downloadservice/")
async def generate_and_downloadv2(request: DocumentRequestV2):
    try:
        # Validar que el id_checklist sea un entero
        if not isinstance(request.id_checklist, int):
            raise HTTPException(status_code=400, detail="El id_checklist debe ser un entero.")

        # Calcular el número de placeholders desde la solicitud
        num_placeholders = len(request.placeholders)

        # Obtener las imágenes en base64 para el checklist proporcionado
        images_base64 = get_service_one(request.id_checklist, num_placeholders)
        logging.info(f"Imágenes obtenidas para id_checklist {request.id_checklist}: {images_base64}")

        if not images_base64:
            raise HTTPException(status_code=404, detail="No se encontraron imágenes para el checklist proporcionado.")

        # Generar el documento Word
        word_stream = generate_word_documentv2(
            request.placeholders,
            images_base64,
            request.logo_base64,
            request.logo_derecho_base64
        )

        # Devolver el documento como respuesta
        return StreamingResponse(
            word_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=EvidenciaFotografica.docx"}
        )

    except HTTPException as e:
        raise e
    except Exception as e:
        logging.error(f"Error inesperado: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generando el documento: {str(e)}")

class DocumentRequest(BaseModel):
    placeholders: Dict[str, str]
    images_base64: List[str]
    logo_base64: str
    logo_derecho_base64: str  # Nuevo parámetro para el logo del lado derecho

def set_header_format(paragraph, text):
    """Establecer formato para el título del encabezado."""
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def generate_word_document(placeholders: Dict[str, str], images_base64: List[str], logo_base64: str, logo_derecho_base64: str) -> BytesIO:
    doc = Document()

    # Crear el encabezado
    section = doc.sections[-1]
    header = section.header
    header_table = header.add_table(rows=1, cols=3, width=Inches(6.5))
    header_table.autofit = True

    # Insertar logo izquierdo
    if logo_base64:
        image_data = base64.b64decode(logo_base64)
        image_stream = BytesIO(image_data)
        left_cell = header_table.cell(0, 0)
        paragraph = left_cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(1.5))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Insertar título centrado
    title_cell = header_table.cell(0, 1)
    title_paragraph = title_cell.paragraphs[0]
    set_header_format(title_paragraph, "FORMATO DE EVIDENCIAS FOTOGRÁFICAS")

    # Insertar logo derecho
    if logo_derecho_base64:
        image_data_right = base64.b64decode(logo_derecho_base64)
        image_stream_right = BytesIO(image_data_right)
        right_cell = header_table.cell(0, 2)
        paragraph_right = right_cell.paragraphs[0]
        run_right = paragraph_right.add_run()
        run_right.add_picture(image_stream_right, width=Inches(1.5))
        paragraph_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Espacio después del encabezado
    doc.add_paragraph()

    # Crear tabla de datos del vehículo
    table_data = doc.add_table(rows=0, cols=2)
    table_data.style = "Table Grid"

    for key, value in placeholders.items():
        row = table_data.add_row()
        row.cells[0].text = key.upper() + ":"
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    # Espacio entre tabla y fotos
    doc.add_paragraph()

    # Crear tabla para imágenes (2 columnas)
    num_images = len(images_base64)
    rows = (num_images + 1) // 2  # Calcular filas necesarias
    table_images = doc.add_table(rows=rows, cols=2)
    table_images.style = "Table Grid"

    for idx, image_base64 in enumerate(images_base64):
        image_data = base64.b64decode(image_base64)
        image_stream = BytesIO(image_data)

        row = idx // 2
        col = idx % 2
        cell = table_images.cell(row, col)

        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(2.5), height=Inches(2.5))

    # Guardar documento en BytesIO
    word_stream = BytesIO()
    doc.save(word_stream)
    word_stream.seek(0)

    return word_stream

@app.post("/generate_and_download/")
async def generate_and_download(request: DocumentRequest):
    try:
        word_stream = generate_word_document(
            request.placeholders, request.images_base64, request.logo_base64, request.logo_derecho_base64
        )
        return StreamingResponse(word_stream,
                                 media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                 headers={"Content-Disposition": "attachment; filename=EvidenciaFotografica.docx"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generando el documento: {str(e)}")

# Definir el modelo para la entrada de datos

class DocumentRequestOrden(BaseModel):
    cliente: str
    telefono: str
    vehiculo: str
    placas: str
    fecha: str
    kilometraje: str
    inventario: Dict[str, str]
    logo_base64: Optional[str]

def generate_word_order(clienteId: int):
    try:
        query = f"exec [Clientes].[ordendeservicio]  @idCliente = {clienteId}"
        with engine.begin() as conn:
            conn.execution_options(autocommit=True)
            roles_df = pd.read_sql(query, conn)

        if roles_df.empty:
            return {"error": "No se encontró información para el cliente."}

        roles_df.fillna("N/A", inplace=True)

        data = {
            "Orden": roles_df['idOrden'].iloc[0],
            "Nombre": roles_df['Nombre'].iloc[0],
            "Facturar a": roles_df['Facturar_a'].iloc[0],
            "Dirección": f"{roles_df['Calle'].iloc[0]}, {roles_df['Colonia'].iloc[0]}, {roles_df['Ciudad'].iloc[0]}, {roles_df['Estado'].iloc[0]}",
            "Teléfono": roles_df['Tel'].iloc[0],
            "Celular": roles_df['Cel'].iloc[0],
            "Email": roles_df['Email'].iloc[0],
            "RFC": roles_df['RFC'].iloc[0]
        }

        vehicle_data = {
            "Marca": roles_df['Marca'].iloc[0],
            "Tipo": roles_df['Tipo'].iloc[0],
            "Modelo": roles_df['Modelo'].iloc[0],
            "Motor": roles_df['Motor'].iloc[0],
            "Color": roles_df['Color'].iloc[0],
            "Kilometraje": roles_df['kms'].iloc[0],
            "No. Serie": roles_df['No_Serie'].iloc[0],
            "Placa": roles_df['Placa'].iloc[0]
        }

        inventory_data = {
            "Espejo Retrovisor": roles_df['Espejo_retrovisor'].iloc[0],
            "Espejo Izquierdo": roles_df['Espejo_izquierdo'].iloc[0],
            "Espejo Derecho": roles_df['Espejo_derecho'].iloc[0],
            "Antena": roles_df['Antena'].iloc[0],
            "Tapones de Ruedas": roles_df['Tapones_ruedas'].iloc[0],
            "Radio": roles_df['Radio'].iloc[0],
            "Encendedor": roles_df['Encendedor'].iloc[0],
            "Gato": roles_df['Gato'].iloc[0],
            "Herramienta": roles_df['Herramienta'].iloc[0],
            "Llanta de Refacción": roles_df['Llanta_refaccion'].iloc[0],
            "Limpiadores": roles_df['Limpiadores'].iloc[0],
            "Pintura Rayada": roles_df['Pintura_rayada'].iloc[0],
            "Cristales Rotos": roles_df['Cristales_rotos'].iloc[0],
            "Golpes": roles_df['Golpes'].iloc[0],
            "Tapetes": roles_df['Tapetes'].iloc[0],
            "Extintor": roles_df['Extintor'].iloc[0],
            "Tapón de Gasolina": roles_df['Tapones_gasolina'].iloc[0],
            "Calaveras Rotas": roles_df['Calaveras_rotas'].iloc[0],
            "Molduras Completas": roles_df['Molduras_completas'].iloc[0],
        }

        doc = Document()

        for section in doc.sections:
            section.top_margin = Pt(30)
            section.bottom_margin = Pt(30)
            section.left_margin = Pt(30)
            section.right_margin = Pt(30)

        # Encabezado con logo en tabla
        header = doc.add_table(rows=1, cols=3)
        row = header.rows[0].cells

        # Reemplaza con tu cadena base64 real del logo
        logo_base64 = "/9j/4AAQSkZJRgABAQEAYABgAAD/2wBDAAMCAgMCAgMDAwMEAwMEBQgFBQQEBQoHBwYIDAoMDAsKCwsNDhIQDQ4RDgsLEBYQERMUFRUVDA8XGBYUGBIUFRT/2wBDAQMEBAUEBQkFBQkUDQsNFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBT/wAARCABWAIwDASIAAhEBAxEB/8QAHwAAAQUBAQEBAQEAAAAAAAAAAAECAwQFBgcICQoL/8QAtRAAAgEDAwIEAwUFBAQAAAF9AQIDAAQRBRIhMUEGE1FhByJxFDKBkaEII0KxwRVS0fAkM2JyggkKFhcYGRolJicoKSo0NTY3ODk6Q0RFRkdISUpTVFVWV1hZWmNkZWZnaGlqc3R1dnd4eXqDhIWGh4iJipKTlJWWl5iZmqKjpKWmp6ipqrKztLW2t7i5usLDxMXGx8jJytLT1NXW19jZ2uHi4+Tl5ufo6erx8vP09fb3+Pn6/8QAHwEAAwEBAQEBAQEBAQAAAAAAAAECAwQFBgcICQoL/8QAtREAAgECBAQDBAcFBAQAAQJ3AAECAxEEBSExBhJBUQdhcRMiMoEIFEKRobHBCSMzUvAVYnLRChYkNOEl8RcYGRomJygpKjU2Nzg5OkNERUZHSElKU1RVVldYWVpjZGVmZ2hpanN0dXZ3eHl6goOEhYaHiImKkpOUlZaXmJmaoqOkpaanqKmqsrO0tba3uLm6wsPExcbHyMnK0tPU1dbX2Nna4uPk5ebn6Onq8vP09fb3+Pn6/9oADAMBAAIRAxEAPwD9QKKKKACiiigAoqC5vrWy2/abqC3Lfd86VUz9Mmof7c0z/oJ2P/gSn+NK6RapzkrpMu0VS/tzTP8AoJ2X/gSn+NH9uaZ/0E7H/wACU/xo5l3K9lU/lf3F2iqX9uaZ/wBBOx/8CU/xo/tzTDgDU7In/r5T/GjmXcPZVP5X9xdoqq2q2MblGv7RWHVWuEBH605dSs3jZ1vbZkU4LCdSBnpk5p7mRYoqv/aVl/z+2v8A3/T/ABqaORJUDxusiHoyEMD+Ip2a3QDqKKKQBRRRQAUUUUAFFFFAHwD/AMFOppY/EfgYJJIg+yXBwjlf4lr4i+1XHaeY/wDbRv8AGv13+Pn7Lnhv9oa+0i613UtSsJNNjeKIWLKAwYgnO4H0ryf/AIdn/D0Z/wCKi8Qf9/Iv/iK+PzDLMViMRKpT2fmf0lwhx1kWTZPRweNT9pG97Rvu77n5wfarjvcTD/tq3+NJ9ruevnzf9/W/xr6P/bE/Zm8Pfs7r4ZbQtR1C/OqGbzftzIduwDG3aB6183W6iSaJGO1WdVY+gJANfL16FXD1fYz39T92ynNMDnOBWYYaN6bvurPTcd9quP8An4mz/wBdW/xqxp95OdQs/wB/N/rk/wCWjf3h71912X7E/wACZ7SCST4phJJI1Zl/tizGCQDiuh8P/wDBPP4Ua4q3uj+NNU1aCGUbpbG8t50DjB2llUgHpx717SyLHxtOSsvM/OaviXw1yyhaV9V8B+Rfxtutb1P49eLNP0+4vp7q41iSCC3hmcs7s+FVRnqTiuz+Imsy+BdB074b6Xqs9wNLk+1a7ew3LkXepMuHUNnlIV/dr6nee9fV37anwM+GH7JPiBvGOg6pq2q/E/xHNNPp9reyxtBpyuCsl0FCg7lyQmT945/hr4EZi2WZizMclmOST3Jr+lOC8l5/9vxC0WkV+b/yP5ArSU6kpR2bZa/tjUGyP7Qux/28P/jX7j/8E+5ZJv2PPhy8sjyu1vcZeRizH/SJepPWvwtOK/c//gnr/wAmc/Dj/r2uP/SiWvW45hGOCpWSXvdPRmKWp9EUUUV+JlhRRRQAUUUUAFFFFABRRRTQHwj/AMFQv9V8P/rdfyWvg+3jW4uIYnzskdUbHXBIB/nX3h/wVC/1PgD63X8lr4Rsf+P61H/TZP8A0IV+c5s7Y9v0P7Q8PP8Akk4f9v8A5s5H9o74Fat+z78T77wzqitLbuq3en3pXC3Vu4yrD3HKkdiK9P8A2Dv2sB+zD8RbtNaeaXwRrMRTUbeEFmilVSYpkX1z8p9Q3tX6RftofsvwftKfBC3WwhRPGOiWy3WlT45k+QF4GPo4H54r8R7yzuNOvJ7S7ge2u7eRopYZBho3U4ZSPUEV/X+S4vDcR5Y8NiV7yST/AEa/rc/jGt/Fl6s774/fGjWPj/8AFTW/GmssyteSbbW1zlba3U4jiH0HX1JJq94L+A2reKPgv44+JtxvtPDvh3yraCQr/wAfd1I6jYvsinLH3ArF+Cvwh1z47fEzRfBfh+Ite6hLiScjK20I5kmf2Vcn34Hev1Z/bM+GOh/Bv/gn7q/g7w9AINL0tbWJGI+eV/NUvK57szZJPvXRmeaU8plh8twtlKTivSN7fiY2Pxt9a/dD/gnr/wAmc/Dj/r2uP/SiWvwwJzX7n/8ABPX/AJM5+HH/AF7XH/pRLXmcd/7lT/xfoyon0RRRRX4gUFFFFABRRRQAUUUUAFFFFAHwj/wVC/1PgD63X8lr4Rsf+P61/wCuqf8AoQr7u/4Khf6v4ff711/Ja+ErP/j9tc/89k/9CFfnObf7+/kf2h4d/wDJJw/7f/Nn7o6SSulWBHXyI/8A0EV+X3/BUX9lI+G9a/4W74Xsv+JZqEiw65awJ/qbg8JOAOz8A++PWv1B0r/kE2Hf9xH/AOgin3tjbalavbXltDd2743QzoHRsHIyCMdRX65lGZVcprxxFL5ruux/GdbWpL1Z8nf8E7f2Vf8AhQnwx/4SXXrQJ448TRJNOsgG6ytTzHb+xPDN74Hat3/gpH/yZ/4xx08y2/8ARor6bJycmvmX/gpF/wAmf+Mf+ult/wCjVrsw+Mq47N6eJrO8pTX5rQxex+INfuh/wT1/5M5+HH/Xtcf+lEtfhh2Nfuf/AME9f+TOfhx/17XH/pTLX6bx1b6lS/xfowifRFFFFfiBQUUUUAFNkljhXdJIkS/3pGCj8zTq4H45fCDwv8bvh7eeHfF1nNe6TG325Yre5kt2EsaNsbchB4yeOlAHcfbLY5xcwEAZP71eB69adHcQTNtjnhkb+6kisfyBr8cvhb8IPD1v/wAE7fir8TUt78eNVnuNIW+a7mI+zfaIMIIycfjjNZOhfDzXvh7rn7PPiKP4eXfw8OsX0O3xPa+IJr+XVS0QIzb7v3W7IOP9rFAH7RNdW8bFWuYFYHBVpVBz6daHureNir3MKMOqtKoI/Amvw50uP4Hy/st+O7/xVq1/H8flvrz7DBJdXgkLeeuzKj92Tt3/AHvxr1342/AHw3Ja/sq61qFpqv8Ab3j+awt/E8k1/OHulMUCkEbvkOCeRg0AfbX7Yn7OerftCR+Fv7G1zStNGmmYv9ukOJA+MFdv0r5yg/4Js+OIL633+LPDysJFfbufcQDngY56V4R+1N4R0nUfjZ8R/Bvhrwp4w1vw78O/Csej6GuhGa4i0y7BEzTXL5yU+eUEknp7Vz3xG+K/iXXfip8HPip4eNxeX3hPwNpWsalbq7Dzfsk5huFPbnv7E15dbLcPiKntZrU+7yzjXOcpwawGEqJU1fSy67n7X2jRWVnBA9xCrQRKj5kAIwAOcnj/AOvU7TRIQGljUtyuXA3fT1r8DNc1Hxb4mk+NfinWlvY73xRoUXiKOFWcmKCfUoWQHHYLgfTFfRPxm+OPhX4weNv2ZbDwR4gm1u70HQ5LbV4baKZDbyi0QEPuUZOUbkZ6V6nSx8LKTk22frSLqBoy4nhMYOC4kUqD6ZzXjv7W/wAK7747/AnX/BOianptjqd88LRy30wEa7HDHOMkcD0r8a7fU/Fvwq/ZVmEst9qHhL4pwERvuYmx1HT7zkZzwHj/ADB/2a9T0WP4M6h8UPjb/wALi1S7tb2HTrZvDS+fdpvufsnIHlZBO7y/vcVtRqyw9SNWG8XdfIR2Un/BI/4pwxq8nirwmiv91muJQD9Ds5r9Hv2Wvhjf/BP4AeFPBusX9le32jwyx3F1ZS5gJaV34Y46Bh17ivyT+Gdz4Z1bUPgJbfHbU9Yt/hyfD+pHzLqa6WPcLqcR7TH82M7BxxjHap9S0nwzrml/HnSPAfjr/hGvhbHrekyaPdatPemyu8CTMW4Kzhj8zfMOi/SvbzHPcbmlONLEyTSd9hJWP23W8t3yFuIXIGcLIpOB360iX1tI4RLmF3P8KyAn+dfjJ8DdQ8N3fi7x14c0m08PaK58KSyjxpoOpanNawZliDQyLMM/P90kLwTxnmvbf2bbrRtH+PXw+0yDT9I1+5llIOqaVfXrPA6o37x1kwvPUjGK+NrYv2NWNO2/mfaZVw7/AGngK2NjUs6fN7tuiV732/U/TSiiivRPjH2CjrweRRRQBUXR9PSze0WwtVs3OWt1gQRsfUrjB/KsDxolzat4bGnW+n7E1KKN0uYVIWIg5EfynY3TBGOlFFAHN6zo+p3mn+K3srLw+NVg1TbaTXVgjItvwWDfJkuRu555xXS3sOp32sTsY9NltLfTRc2f2iANJFdcjOSOEGAeOe1FFAEFs1/b+LNQgWOxSK40xZz5cKqZJcEFpDty3zepxjtVa10+aLxdp1kbLSUsZNNzcxR2aDdkfOB8vQsVwM4xnIzRRQBautPvP+Ei1JLew0ZdKOk+XDvgHnGYc4cbcGLBA257dKqeFdJlvLXw1qJs9Js3mErXq2tlGDIhUhFVtoIx36ZFFFAGPcQ65ceDdbjt7bQZbuxvttpDdWam2TB/eDaE4BB4IBPJzV5dJuZNc8LxXthoc73Fn5moN9hT946gbypK5AwRt5HQ5oooAl8Q6LrJ0SzSEaNdNb6mwIvLFGRLLfjyo1C4VtuOam17TzZwR2cFhpUVtNrEaiNbSMp5G0E5Qrgv1Ge2etFFAHVQ+HtJtVmSDSrGBJeHWO2RQ49Dgc/jUlto+n2cgkt7C1t5AMB4YVVh+IFFFKyLU5RXKnoW6KKKZB//2Q=="
        image_stream = BytesIO(base64.b64decode(logo_base64))
        row[0].paragraphs[0].add_run().add_picture(image_stream, width=Inches(1.2))

        row[1].text = (
            "ORDEN DE SERVICIO\nServicio Automotriz Trotamundos\n"
            "29 Guerrero y Bravo #422\nCol. Héroe de Nacozari, C.P. 87030\n"
            "Tel: (834) 285 2869 / (834) 285 2872\nR.F.C. GACM040320DD9"
        )
        for p in row[1].paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in p.runs:
                run.font.size = Pt(8)

        row[2].text = f"ORDEN\nNo. {data['Orden']}"
        row[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        def set_cell_width(cell, width_cm):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(int(width_cm * 567)))  # 1 cm ~ 567 twips
            tcW.set(qn('w:type'), 'dxa')

        def set_table_borders(table, border_size=24):
            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is None:
                tblBorders = OxmlElement('w:tblBorders')
                tblPr.append(tblBorders)

            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = tblBorders.find(qn(f'w:{border_name}'))
                if border is None:
                    border = OxmlElement(f'w:{border_name}')
                    tblBorders.append(border)
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), str(border_size))  # grosor del borde
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # negro

        def set_font(run, font_name='Times New Roman', font_size_pt=10, bold=True):
            run.font.name = font_name
            run.font.size = Pt(font_size_pt)
            run.font.bold = bold
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), font_name)

        def add_table(data_dict, title):
            # Crear tabla con 1 fila más para encabezado y 2 columnas
            table = doc.add_table(rows=len(data_dict) + 1, cols=2)
            table.style = 'Table Grid'

            # Ajustar bordes con mayor grosor
            set_table_borders(table, border_size=24)  # 24 = 1.2 pt

            # Ajustar ancho columnas (ejemplo: 5cm y 7cm)
            for row in table.rows:
                set_cell_width(row.cells[0], 5)
                set_cell_width(row.cells[1], 7)

            # Encabezado - fila 0, celdas 0 y 1
            header_cells = table.rows[0].cells
            header_cells[0].text = title
            header_cells[1].text = ""

            # Estilo para encabezado (negrita, texto blanco, fondo gris)
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_font(run, font_name='Times New Roman', font_size_pt=11, bold=True)
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), '808080')  # gris oscuro
                cell._tc.get_or_add_tcPr().append(shading_elm)

            # Agregar filas con datos desde la fila 1 en adelante
            for i, (key, val) in enumerate(data_dict.items(), start=1):
                key_cell = table.cell(i, 0)
                val_cell = table.cell(i, 1)

                key_cell.text = str(key)
                val_cell.text = str(val)

                # Estilo para texto en filas de datos
                for cell in (key_cell, val_cell):
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            set_font(run, font_name='Times New Roman', font_size_pt=9, bold=True)

                # Ajustar altura fila datos (más pequeña)
                tr = key_cell._tc.getparent()
                trPr = tr.get_or_add_trPr()
                trHeight = OxmlElement('w:trHeight')
                trHeight.set(qn('w:val'), "240")  # aprox 12 pts para fila más pequeña
                trHeight.set(qn('w:hRule'), 'exact')
                trPr.append(trHeight)

            return table


        def add_section(title: str, content: dict):
            # En vez de un párrafo para título, el título está en encabezado de tabla
            add_table(content, title)


        # Agregar secciones
        add_section("Información del Cliente", data)
        add_section("Detalles del Vehículo", vehicle_data)
        add_section("Inventario del Vehículo", inventory_data)

        tabla_obs = doc.add_table(rows=2, cols=1)
        tabla_obs.style = 'Table Grid'  # Asegura que tenga bordes visibles

# Fila 0: encabezado
        set_table_borders(tabla_obs, border_size=24)  # grosor 1.2 pt como en las otras tablas

        # Celda de encabezado
        celda_encabezado = tabla_obs.cell(0, 0)
        celda_encabezado.text = "Observaciones"

        for parrafo in celda_encabezado.paragraphs:
            for run in parrafo.runs:
                set_font(run, font_name='Times New Roman', font_size_pt=11, bold=True)
                run.font.color.rgb = RGBColor(255, 255, 255)  # blanco
            parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Fondo gris del encabezado
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '808080')  # gris oscuro
        celda_encabezado._tc.get_or_add_tcPr().append(shading_elm)

        # Fila 1: celda vacía para observaciones
        celda_obs = tabla_obs.cell(1, 0)
        celda_obs.text = ""  # vacía, espacio para escribir observaciones

        # Ajustar altura de la fila de observaciones
        tr = celda_obs._tc.getparent()
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), "360")  # 18 puntos de alto
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

        # Ajustar fuente de la celda de observaciones
        for parrafo in celda_obs.paragraphs:
            for run in parrafo.runs:
                set_font(run, font_name='Times New Roman', font_size_pt=9, bold=True)
            parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        firmas = doc.add_paragraph()
        firmas.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        tabla_firmas = doc.add_table(rows=2, cols=2)
        tabla_firmas.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        tabla_firmas.autofit = True

        tabla_firmas.cell(0, 0).text = "__________________________"
        tabla_firmas.cell(0, 1).text = "__________________________"

        tabla_firmas.cell(1, 0).text = "Firma del Proveedor"
        tabla_firmas.cell(1, 1).text = "Firma del Cliente"

        # Estilo de fuente en la tabla de firmas
        for row in tabla_firmas.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.bold = True
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



        # Segunda página
        doc.add_page_break()
        contrato_text = """ADEMAS DE LOS ELEMENTOS CONTENIDOS EN EL ANVERSO DEL PRESENTE CONTRATO DE PRESTACIÓN DE SERVICIOS DE REPARACIÓN DE VEHÍCULOS, LAS PARTES SE SUJETAN A LAS SIGUIENTES: Y MANT

        CLAUSULAS

        PRIMERA: EL PRESTADOR DEL SERVICIO realizará todas las operaciones y composturas descritas en el anverso del presente contrato, solicitadas por EL CONSUMIDOR suscribe el presente contrato, a las que se someterà el vehiculo para obtener condiciones de funcionamiento de acuerdo al estado de este Asimismo EL PRESTADOR DEL SERVICIO no condicionară la prestación de los servicios de reparación y/o mantenimiento de vehículos a la adquisición a renta de otros productos o servicios en el establecimiento o en otro taller o agencia predeterminada que

        May

        SEGUNDA: El precio total de los servicios contratados se establece en el presupuesto que forma parte del presente y se describe en el anverso del presente contrato, el cual sera pagado por EL CONSUMIDOR, de la siguiente forma: Al momento de celebrar el presente contrato por concepto de anticipo la cantidad que se indica en el anverso del presente contrato y el resto en la fecha de entrega del vehículo, Todo pago efectuado por EL CONSUMIDOR deberá realizarse en el establecimiento de EL PRESTADOR DEL SERVICIO, al contado y en moneda nacional o cualquier moneda extranjera aceptada por EL PRESTADOR DEL SERVICIO este último deberá estar conforme a la Ley Monetaria de los Estados Unidos Mexicanos.

        El pago será en efectivo, salvo que las partes acuerden o acepten otra forma distinta, como pudiese ser en cheque tarjeta de crédito o deposito bancario.

        TERCERA: EL PRESTADOR DEL SERVICIO pondrá a disposición de EL CONSUMIDOR los precios de los servicios, mano de obra, refacciones y materiales a usar en las reparaciones ofrecidas. Asimismo, previo a la realización del servicio EL PRESTADOR DEL SERVICIO presentará a EL CONSUMIDOR el presupuesto al que se refiere la cláusula Segunda del presente contrato. Una vez aprobado el presupuesto por EL CONSUMIDOR, EL PRESTADOR DE SERVICIO procederá a efectuar el servicio solicitado. Los incrementos que resulten durante la reparación por costos no previsibles en rubros específicos que su cotización este fuera de control de EL PRESTADOR DEL SERVICIO, deberán ser autorizados por EL CONSUMIDOR en forma escrita, siempre y cuando estos excedan al 20% del presupuesto. Si el incremento citado es inferior lo podrán autorizar telefónicamente. El tiempo, que en su caso, transcurra para requisitar esta condición se modificará la fecha de entrega, en la misma proporción.

        CUARTA: La entrega del automóvil será en la fecha contemplada en el anverso del presente contrato. Para el caso de que EL CONSUMIDOR, sea el que proporcione las refacciones la fecha de entrega sera

        QUINTA: EL PRESTADOR DEL SERVICIO exclusivamente utilizará para los servicios objeto de este contrato, partes, refacciones u otros materiales nuevos y apropiados para el vehiculo salvo que EL CONSUMIDOR autorice expresamente que se usen otras. SI SI EL EL PRESTADOR PRES DEL SERVICIO lo autoriza, EL CONSUMIDOR suministra las partes, refacciones o materiales necesarios para la reparación y/o mantenimiento del vehículo. En ambos casos, la autorización respectiva se hará constar en el anverso del presente contrato.

        SEXTA: EL PRESTADOR DEL SERVICIO hará entrega de las refacciones, partes o plezas sustituidas en la reparación y/o mantenimiento del vehículo al momento de entrega de éste, salvo en los siguientes casos: a) cuando EL CONSUMIDOR, exprese lo contrario, b) as partes, refacciones o piezas sean cambiadas en uso de garantía, c) se trate de residuos considerados peligrosos de acuerdo con las disposiciones legales aplicables.

        SEPTIMA: reparaciones a que se refiere el presupuesto aceptado por EL CONSUMIDOR Tienen una garantia de 60 días contados a partir de la fecha de entrega del vehiculo ya reparado en mano de obra, y refacciones la especificada por el fabricante, siempre y cuando no se manifieste mal uso, negligencia o descuido, lo anterior de conformidad a lo establecido con el articulo 77 de la Ley Federal de Protección al Consumidor. Si el vehiculo es intervenido por un tercero, "EL PRESTADOR DEL SERVICIO" no será responsable y la garantia quedara sin efecto. Las Jedara sin efecto. Las reclamaciones por garantia se harán en el establecimiento de EL PRESTADOR DEL SERVICIO para lo cual EL CONSUMIDOR deberá presentar damaciones por garantia se harán en el su vehiculo en dicho establecimiento. Las reparaciones nes efectuadas por EL PRESTADOR DEL SERVICIO en cumplimiento a la garantia del servicio, serán sin cargo alguno para EL CONSUMIDOR salvo aquellos trabajos que no deriven de las reparaciones aceptadas en el presupuesto. No se computara dentro del plazo de garantía, el tiempo que lleve la reparación yio mantenimiento del vehiculo para el cumplimiento de la misma. Los gastos en que incurra EL CONSUMIDOR para hacer valida la garantia en un domicilio diverso al de EL PRESTADOR DEL SERVICIO deberán ser cubiertos por este. STADOR DEL SERVICio pa

        OCTAVA: EL CONSUMIDOR autoriza el uso del vehiculo en zonas aledañas con un radio de 5 Km al área del establecimiento a efectos de pruebas o verificación de las reparaciones a efectuar o efectuadas EL PRESTADOR DEL SERVICIO no podrá utilizar el vehiculo para uso personal, fines propios o de terceros.

        NOVENA: EL PRESTADOR DEL SERVICIO se hace responsable por los daños causados al vehiculo de EL CONSUMIDOR, como consecuencia de los recorridos de prueba por parte del personal de EL PRESTADOR DEL SERVICIO. El riesgo en un recorrido de prueba es por cuenta de EL CONSUMIDOR cuando él mismo solicite que sera él o un representante suyo quien gule el vehiculo, Asimismo, EL PRESTADOR DEL SERVICIO se hace responsable por las descomposturas, daños perdidas parciales o totales, imputables a él o a sus empleados, que sufran los vehículos, el equipo y aditamentos que EL CONSUMIDOR haya notificado al momento de la recepción del vehiculo, mientras encuentren bajo su respons onsabilidad para llevar a cabo la reparación y/o mantenimiento solicitado, asi como para hacer efectiva la garantia otorgada, Para tal efecto EL PRESTADOR DEL SERVICIO SI )NO() cuenta con un seguro suficiente para cubrir dichas eventualidades, cuyo número de póliza es con la compañía EL PRESTADOR DEL SERVICIO no se hace responsable por la pérdida de objetos dejados en el interior del vehiculo, aún con la cajuela cerrada, salvo que estos hayan sido notificados y puestos bajo su resguardo al momento de la recepción del vehículo.

        DECIMA: EL PRESTADOR DEL SERVICIO se obliga a expedir la factura o comprobante de pago por los trabajos efectuados, en la que se especificará los precios por mano de obra, refacciones, materiales y accesorios empleados, conforme al articulo 62 de la Ley Federal de Protección al Consumidor

        DECIMA PRIMERA: Se establece como pena convencional por el incumplimiento de cualquiera de las partes a las obligaciones contraldas en el presente contrato, el 15% del precio total de la operación.

        DECIMA SEGUNDA: En caso de que el vehículo no sea recogido por EL CONSUMIDOR en un plazo de 48 horas a partir de la fecha señalada para la entrega, pagará por concepto de deposito un salario mínimo vigente en el lugar que se celebre el presente contrato, por cada 24 hrs. que transcurran.

        DECIMA TERCERA: EL CONSUMIDOR puede desistirse en cualquier momento de la nto de la contratación del servicio de reparación y/o mantenimiento del vehículo, en cuyo caso deberá cubrir en lugar del precio contratado el importe de los trabajos realizados hasta el retiro del vehículo, incluidas las partes, refacciones u otros materiales utilizados. DECIMA CUARTA: EL PRESTADOR DEL SERVICIO es responsable ante EL CONSUMIDOR por el incumplimiento de los servicios contratados, aún y cuando subcontrate con

        lerceros dicha prestación.

        DECIMA QUINTA: Cuando se preste el servicio a presentación vicio a domicilio, el personal de EL PRESTADOR DEL SERVICIO debe identificarse plenamente ante EL CONSUMIDOR, mediante la del documento que lo acredite para este proposito. En caso de que dicho servicio tenga un costo, este se indicara en el anverso del presente contrato.

        DECIMA SEXTA: EL CONSUMIDOR libera a EL PRESTADOR DEL SERVICIO de cualquier responsabilidad que hubiere surgido o pudiese surgir con relación al origen, propiedad,

        posesión o cualquier otro derecho inherente al vehículo o partes o componentes del mismo.

        DECIMA SEPTIMA: EL PRESTADOR DEL SERVICIO se obliga a observar en lo relativo a información y publicidad, promociones y ofertas a lo dispuesto en los capítulos II y IV de la Ley Federal de Protección al Consumidor.

        DECIMA NOVENA: La Procuraduría Federal del Consumidor es competente para conocer en la via administrativa, para resolver cualquier controversia que se suscite sobre la interpretación o cumplimiento del presente Contrato. Sin perjuicio de lo Sin perjuicio de lo anterior, las partes se someten a la jurisdicción de los tribunales competen competentes de la Ciudad de Victoria, Tam.. renunciando expresamente a cualquier otra jurisdicción que pudiera corresponderles por razón de sus domicilios presentes o futuros, o por cualquier otra razón.

        EL PRESTADOR DEL SERVICIO

        EL CONSUMIDOR

        QUEJAS Y RECLAMACIONES A LOS TELEFONOS (834) 285 2869 y (834) 285 2872 Domicilio: 29 Guerrero y Bravo #422. Héroe de Nacozari. C.P. 87030 Ciudad Victoria, Tam.

        TROTAMUNDOSS

        TELEFONOS (834) 285 2869 y (834) 285 2872

        Domicilio: 29 Guerrero y Bravo #422. Héroe de Nacozari. C.P. 87030 Ciudad Victoria, Tam
        """

        # Divide en párrafos por cláusula
        for clausula in contrato_text.split("\n\n"):
            parrafo = doc.add_paragraph(clausula.strip())
            parrafo.paragraph_format.line_spacing = 1.0
            for run in parrafo.runs:
                run.font.size = Pt(5.5)  # Más pequeño para compactar

        # REDUCIR MÁRGENES
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # GUARDAR EN MEMORIA
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
        

    except Exception as e:
        return {"error": str(e)}

# Endpoint para descarga
@app.get("/generate_and_download_orden")
async def generate_and_download(clienteId: int):
    file_stream = generate_word_order(clienteId)
    
    if isinstance(file_stream, dict):
        return file_stream

    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=orden_servicio.docx"}
    )


@app.post(
    path="/api/seguridad/iniciarsesion",
    name='Inicio de sesion',
    tags=['Seguridad'],
    description='Método para iniciar sesión',
    response_model=Token
)
async def login(payload: DatosLogin):
    _negocios = Negocios()
    user = await _negocios.getusuario(payload)
    if not user:
        return JSONResponse(
            status_code=401,
            content={"Id_Resultado": 0, "Respuesta": "Datos de acceso incorrectos"}
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = await utilsclass.create_access_token(
        data={"sub": payload.usuario, "idUsuario": user["IdUsuario"], "idRol": user["Rol"]},
        expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}



# app.get(
#         path="/api/cliente",
#         name='Obtener clien',
#         tags=['Cliente'],
#         description='Método para obtener la informacion de un cliente',
#         response_model=GetCliente
# )
# def getcliente(IdCliente = int):
#     query = f"exec Clientes.clienteinsupdel @Accion = 4,@IdCliente ={IdCliente}"
#     roles_df = pd.read_sql(query, engine)
#     resultado = roles_df.to_dict(orient="records")
#     return JSONResponse(status_code=200,content=resultado)
# @app.get(
#         path="/api/clientes",
#         name='Obtener clientes',
#         tags=['Cliente'],
#         description='Método para obtener la informacion del cliente',
#         response_model=List[GetCliente]
# )
# def getclientes(busqueda = ""):
#     query = f"exec Clientes.clienteinsupdel @Accion = 2,@ParametroBusqueda = '{busqueda}' "
#     roles_df = pd.read_sql(query, engine)
#     resultado = roles_df.to_dict(orient="records")
#     return JSONResponse(status_code=200,content=resultado)

@app.get(
    path="/api/cliente",
    name='Obtener cliente',
    tags=['Cliente'],
    description='Método para obtener la informacion de un cliente',
    response_model=List[GetCliente]
)
def getcliente(idCliente=0):
    query = f"exec Clientes.clienteinsupdel @Accion = 4,@IdCliente ={idCliente}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    resultado = resultado[0]  # Solo toma el primer registro, lo cual podría ser un problema si no se encuentra el cliente
    return JSONResponse(status_code=200, content=resultado)


@app.get(
    path="/api/clientes",
    name='Obtener clientes',
    tags=['Cliente'],
    description='Método para obtener la informacion de todos los clientes',
    response_model=List[GetCliente]
)
def getclientes(busqueda=""):
    query = f"exec Clientes.clienteinsupdel @Accion = 2,@ParametroBusqueda = '{busqueda}' "
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200, content=resultado)
@app.post(
    path="/api/cliente",
    name='Guardar cliente',
    tags=['Cliente'],
    description='Método para guardar la información del cliente',
    response_model=ResponseModel
)
def saveCliente(payload: SaveCliente):
    # Convertir valores None a cadenas vacías
    payload_dict = payload.model_dump()
    for key, value in payload_dict.items():
        if value is None:
            payload_dict[key] = ""

    query = f"EXEC clientes.clienteinsupdel \
    @Nombre='{payload_dict['Nombre']}', \
    @Calle='{payload_dict['Calle']}',  \
    @Colonia='{payload_dict['Colonia']}', \
    @Ciudad='{payload_dict['Ciudad']}', \
    @Estado='{payload_dict['Estado']}', \
    @Tel='{payload_dict['Tel']}', \
    @Cel='{payload_dict['Cel']}', \
    @Email='{payload_dict['Email']}', \
    @RFC='{payload_dict['RFC']}', \
    @No_int='{payload_dict['No_int']}', \
    @Facturar_a='{payload_dict['Facturar_a']}', \
    @IdUsuarioEmpleado='{payload_dict['Id_empleado']}', @Accion = 1"

    print(query)

    with engine.begin() as conn:
        conn.execution_options(autocommit=True)
        roles_df = pd.read_sql(query, conn)

    dumpp = ResponseModel(id_resultado=1, respuesta="El cliente se guardó de manera correcta")
    dict_response = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict_response)



#################################################################################################################
@app.get(
    path="/api/orderservice",
    name='Obtener orden de servicio',
    tags=['Orden'],
    description='Método para obtener la información de una orden de servicio',
    response_model=GetOrden
)
def getorder(clienteId=0):
    query = f"exec [Clientes].[ordendeservicio]  @idCliente = {clienteId}"
    cliente_df = pd.read_sql(query, engine)
    resultado = cliente_df.to_dict(orient="records")

    # Verifica si el resultado no está vacío antes de devolver
    if resultado:
        return JSONResponse(status_code=200, content=resultado[0])
    else:
        return JSONResponse(status_code=404, content={"message": "Orden no encontrado."})
@app.get(
    path="/api/orderservices",
    name='Obtener ordenes de servicio',
    tags=['Orden'],
    description='Método para obtener la información de todos las ordenes de servicio',
    response_model=List[GetOrden]
)
def getorders():
    query = f"exec [dbo].[OrdenDeServicioAll]"
    orders_df = pd.read_sql(query, engine)
    resultado = orders_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)




@app.get(
    path="/api/vehiculo",
    name='Obtener vehiculo',
    tags=['Vehiculo'],
    description='Método para obtener la información de un vehículo',
    response_model=GetVehiculo
)
def getvehiculo(idVehiculo=0):
    query = f"exec [dbo].[ObtenerVehiculo] @IdVehiculo = {idVehiculo}"
    vehiculo_df = pd.read_sql(query, engine)
    resultado = vehiculo_df.to_dict(orient="records")

    # Verifica si el resultado no está vacío antes de devolver
    if resultado:
        return JSONResponse(status_code=200, content=resultado[0])
    else:
        return JSONResponse(status_code=404, content={"message": "Vehículo no encontrado."})

@app.get(
    path="/api/vehiculos",
    name='Obtener vehiculos',
    tags=['Vehiculo'],
    description='Método para obtener la información de todos los vehículos',
    response_model=List[GetVehiculo]
)
def getvehiculos(parametro=""):
    query = f"exec [dbo].[ObtenerVehiculo] @ParametroBusqueda = '{parametro}'"
    vehiculos_df = pd.read_sql(query, engine)
    resultado = vehiculos_df.to_dict(orient="records")

    # Devuelve la lista de vehículos
    return JSONResponse(status_code=200, content=resultado)
#################################################################################################################
# @app.get(
#         path="/api/vehiculo",
#         name='Obtener vehiculo',
#         tags=['Vehiculo'],
#         description='Método para obtener la informacion de todos los vehiculos',
#         response_model=List[GetVehiculo]
# )
# def getvehiculos(idVehiculo = 0):
#     query = f"exec [dbo].[ObtenerVehiculo] @IdVehiculo = {idVehiculo}"
#     roles_df = pd.read_sql(query, engine)
#     resultado = roles_df.to_dict(orient="records")
#     return JSONResponse(status_code=200,content=resultado[0])


# @app.get(
#     path="/api/vehiculos",
#         name='Obtener vehiculos',
#         tags=['Vehiculo'],
#         description='Método para obtener la informacion de un vehiculo',
#         response_model=GetVehiculo
# )
# def getvehiculos(parametro = ""):
#     query = f"exec [dbo].[ObtenerVehiculo] @ParametroBusqueda = '{parametro}'"
#     roles_df = pd.read_sql(query, engine)
#     resultado = roles_df.to_dict(orient="records")
#     return JSONResponse(status_code=200,content=resultado)

@app.post(
        path="/api/vehiculo",
        name='Guarda vehiculo',
        tags=['Vehiculo'],
        description='Método para guardar la informacion de los vehiculos del cliente}',
    response_model=ResponseModel
)
def guardarVehiculo(payload: saveVehiculo):
    try:
        # Convertir listas de fotos a cadenas de Base64 separadas por comas
        fotos = {
            "MotorVehiculo_foto": ",".join(payload.MotorVehiculo_foto),
            "Acumulador_foto": ",".join(payload.Acumulador_foto),
            "Espejo_retrovisor_foto": ",".join(payload.Espejo_retrovisor_foto),
            "Espejo_izquierdo_foto": ",".join(payload.Espejo_izquierdo_foto),
            "Espejo_derecho_foto": ",".join(payload.Espejo_derecho_foto),
            "Antena_foto": ",".join(payload.Antena_foto),
            "Tapones_ruedas_foto": ",".join(payload.Tapones_ruedas_foto),
            "Radio_foto": ",".join(payload.Radio_foto),
            "Encendedor_foto": ",".join(payload.Encendedor_foto),
            "Gato_foto": ",".join(payload.Gato_foto),
            "Herramienta_foto": ",".join(payload.Herramienta_foto),
            "Llanta_refaccion_foto": ",".join(payload.Llanta_refaccion_foto),
            "Limpiadores_foto": ",".join(payload.Limpiadores_foto),
            "Pintura_rayada_foto": ",".join(payload.Pintura_rayada_foto),
            "Cristales_rotos_foto": ",".join(payload.Cristales_rotos_foto),
            "Golpes_foto": ",".join(payload.Golpes_foto),
            "Tapetes_foto": ",".join(payload.Tapetes_foto),
            "Extintor_foto": ",".join(payload.Extintor_foto),
            "Tapones_gasolina_foto": ",".join(payload.Tapones_gasolina_foto),
            "Calaveras_rotas_foto": ",".join(payload.Calaveras_rotas_foto),
            "Molduras_completas_foto": ",".join(payload.Molduras_completas_foto),
            "Panel_instrumentos_foto": ",".join(payload.Panel_instrumentos_foto),
            "Lado_izquierdo_foto": ",".join(payload.Lado_izquierdo_foto),
            "Lado_izquierdo_inf_foto": ",".join(payload.Lado_izquierdo_inf_foto),
            "Lado_derecho_foto": ",".join(payload.Lado_derecho_foto),
            "Lado_derecho_inf_foto": ",".join(payload.Lado_derecho_inf_foto),
            "Tablero_foto": ",".join(payload.Tablero_foto),
            "Guantera_foto": ",".join(payload.Guantera_foto),
            "Consola_foto": ",".join(payload.Consola_foto),
            "LadoFrontal_foto": ",".join(payload.LadoFrontal_foto),
            "LadoTrasero_foto": ",".join(payload.LadoTrasero_foto),
            "Cajuela_foto": ",".join(payload.Cajuela_foto),
            "NumeroEconomico_foto": ",".join(payload.NumeroEconomico_foto),




            "MotorVehiculo_video": ",".join(payload.MotorVehiculo_video),
            "Acumulador_video": ",".join(payload.Acumulador_video),
            "Espejo_retrovisor_video": ",".join(payload.Espejo_retrovisor_video),
            "Espejo_izquierdo_video": ",".join(payload.Espejo_izquierdo_video),
            "Espejo_derecho_video": ",".join(payload.Espejo_derecho_video),
            "Antena_video": ",".join(payload.Antena_video),
            "Tapones_ruedas_video": ",".join(payload.Tapones_ruedas_video),
            "Radio_video": ",".join(payload.Radio_video),
            "Encendedor_video": ",".join(payload.Encendedor_video),
            "Gato_video": ",".join(payload.Gato_video),
            "Herramienta_video": ",".join(payload.Herramienta_video),
            "Llanta_refaccion_video": ",".join(payload.Llanta_refaccion_video),
            "Limpiadores_video": ",".join(payload.Limpiadores_video),
            "Pintura_rayada_video": ",".join(payload.Pintura_rayada_video),
            "Cristales_rotos_video": ",".join(payload.Cristales_rotos_video),
            "Golpes_video": ",".join(payload.Golpes_video),
            "Tapetes_video": ",".join(payload.Tapetes_video),
            "Extintor_video": ",".join(payload.Extintor_video),
            "Tapones_gasolina_video": ",".join(payload.Tapones_gasolina_video),
            "Calaveras_rotas_video": ",".join(payload.Calaveras_rotas_video),
            "Molduras_completas_video": ",".join(payload.Molduras_completas_video),
            "Panel_instrumentos_video": ",".join(payload.Panel_instrumentos_video),
            "Lado_izquierdo_video": ",".join(payload.Lado_izquierdo_video),
            "Lado_izquierdo_inf_video": ",".join(payload.Lado_izquierdo_inf_video),
            "Lado_derecho_video": ",".join(payload.Lado_derecho_video),
            "Lado_derecho_inf_video": ",".join(payload.Lado_derecho_inf_video),
            "Tablero_video": ",".join(payload.Tablero_video),
            "Guantera_video": ",".join(payload.Guantera_video),
            "Consola_video": ",".join(payload.Consola_video),
            "LadoFrontal_video": ",".join(payload.LadoFrontal_video),
            "LadoTrasero_video": ",".join(payload.LadoTrasero_video),
            "Cajuela_video": ",".join(payload.Cajuela_video),
            "NumeroEconomico_video": ",".join(payload.NumeroEconomico_video)

        }

        # Crear el diccionario de parámetros sin conflicto
        parametros = payload.dict(exclude=fotos.keys())
        parametros.update(fotos)

        query = text("""
            exec dbo.InsertarVehiculo
                @Id_Cliente = :Id_Cliente,
                @Id_Empleado = :Id_Empleado,
                @Marca = :Marca,
                @Modelo = :Modelo,
                @Color = :Color,
                @No_serie = :No_serie,
                @Placa = :Placa,
                @Tipo = :Tipo,
                @Motor = :Motor,
                @Kms = :Kms,
                @MotorVehiculo = :MotorVehiculo,
                @Acumulador = :Acumulador,
                @Espejo_retrovisor = :Espejo_retrovisor,
                @Espejo_izquierdo = :Espejo_izquierdo,
                @Espejo_derecho = :Espejo_derecho,
                @Antena = :Antena,
                @Tapones_ruedas = :Tapones_ruedas,
                @Radio = :Radio,
                @Encendedor = :Encendedor,
                @Gato = :Gato,
                @Herramienta = :Herramienta,
                @Llanta_refaccion = :Llanta_refaccion,
                @Limpiadores = :Limpiadores,
                @Pintura_rayada = :Pintura_rayada,
                @Cristales_rotos = :Cristales_rotos,
                @Golpes = :Golpes,
                @Tapetes = :Tapetes,
                @Extintor = :Extintor,
                @Tapones_gasolina = :Tapones_gasolina,
                @Calaveras_rotas = :Calaveras_rotas,
                @Molduras_completas = :Molduras_completas,
                @Panel_instrumentos = :Panel_instrumentos,
                @Lado_izquierdo = :Lado_izquierdo,
                @Lado_izquierdo_inf = :Lado_izquierdo_inf,
                @Lado_derecho = :Lado_derecho,
                @Lado_derecho_inf = :Lado_derecho_inf,
                @Tablero = :Tablero,
                @Guantera = :Guantera,
                @Consola = :Consola,
                @LadoFrontal = :LadoFrontal,
                @LadoTrasero = :LadoTrasero,
                @Cajuela = :Cajuela,
                @NumeroEconomico = :NumeroEconomico,
                @MotorVehiculo_foto = :MotorVehiculo_foto,
                @Acumulador_foto = :Acumulador_foto,
                @Espejo_retrovisor_foto = :Espejo_retrovisor_foto,
                @Espejo_izquierdo_foto = :Espejo_izquierdo_foto,
                @Espejo_derecho_foto = :Espejo_derecho_foto,
                @Antena_foto = :Antena_foto,
                @Tapones_ruedas_foto = :Tapones_ruedas_foto,
                @Radio_foto = :Radio_foto,
                @Encendedor_foto = :Encendedor_foto,
                @Gato_foto = :Gato_foto,
                @Herramienta_foto = :Herramienta_foto,
                @Llanta_refaccion_foto = :Llanta_refaccion_foto,
                @Limpiadores_foto = :Limpiadores_foto,
                @Pintura_rayada_foto = :Pintura_rayada_foto,
                @Cristales_rotos_foto = :Cristales_rotos_foto,
                @Golpes_foto = :Golpes_foto,
                @Tapetes_foto = :Tapetes_foto,
                @Extintor_foto = :Extintor_foto,
                @Tapones_gasolina_foto = :Tapones_gasolina_foto,
                @Calaveras_rotas_foto = :Calaveras_rotas_foto,
                @Molduras_completas_foto = :Molduras_completas_foto,
                @Panel_instrumentos_foto = :Panel_instrumentos_foto,
                @Lado_izquierdo_foto = :Lado_izquierdo_foto,
                @Lado_izquierdo_inf_foto = :Lado_izquierdo_inf_foto,
                @Lado_derecho_foto = :Lado_derecho_foto,
                @Lado_derecho_inf_foto = :Lado_derecho_inf_foto,
                @Tablero_foto = :Tablero_foto,
                @Guantera_foto = :Guantera_foto,
                @Consola_foto = :Consola_foto,
                @LadoFrontal_foto = :LadoFrontal_foto,
                @LadoTrasero_foto = :LadoTrasero_foto,
                @Cajuela_foto = :Cajuela_foto,
                @NumeroEconomico_foto = :NumeroEconomico_foto,
                @MotorVehiculo_video = :MotorVehiculo_video,
                @Acumulador_video = :Acumulador_video,
                @Espejo_retrovisor_video = :Espejo_retrovisor_video,
                @Espejo_izquierdo_video = :Espejo_izquierdo_video,
                @Espejo_derecho_video = :Espejo_derecho_video,
                @Antena_video = :Antena_video,
                @Tapones_ruedas_video = :Tapones_ruedas_video,
                @Radio_video = :Radio_video,
                @Encendedor_video = :Encendedor_video,
                @Gato_video = :Gato_video,
                @Herramienta_video = :Herramienta_video,
                @Llanta_refaccion_video = :Llanta_refaccion_video,
                @Limpiadores_video = :Limpiadores_video,
                @Pintura_rayada_video = :Pintura_rayada_video,
                @Cristales_rotos_video = :Cristales_rotos_video,
                @Golpes_video = :Golpes_video,
                @Tapetes_video = :Tapetes_video,
                @Extintor_video = :Extintor_video,
                @Tapones_gasolina_video = :Tapones_gasolina_video,
                @Calaveras_rotas_video = :Calaveras_rotas_video,
                @Molduras_completas_video = :Molduras_completas_video,
                @Panel_instrumentos_video = :Panel_instrumentos_video,
                @Lado_izquierdo_video = :Lado_izquierdo_video,
                @Lado_izquierdo_inf_video = :Lado_izquierdo_inf_video,
                @Lado_derecho_video = :Lado_derecho_video,
                @Lado_derecho_inf_video = :Lado_derecho_inf_video,
                @Tablero_video = :Tablero_video,
                @Guantera_video = :Guantera_video,
                @Consola_video = :Consola_video,
                @LadoFrontal_video = :LadoFrontal_video,
                @LadoTrasero_video = :LadoTrasero_video,
                @Cajuela_video = :Cajuela_video,
                @NumeroEconomico_video = :NumeroEconomico_video,
                @IdFlotilla = :IdFlotilla,
                @IdOrdenServicio = :IdOrdenServicio,
                @Activo = :Activo


        """)

        # Ejecutar la consulta pasando `parametros` como un solo diccionario
        with engine.begin() as conn:
            conn.execute(query, parametros)

        # Respuesta de éxito
        return JSONResponse(status_code=200, content={
            "id_resultado": 1,
            "respuesta": "Se guardó la información del vehículo de manera correcta",
            "detalles": parametros
        })

    except Exception as e:
        # Respuesta de error
        raise HTTPException(status_code=500, detail=f"Error al guardar el vehículo: {str(e)}")
@app.put(
    path="/api/vehiculoporid",
    name="Actualizar estado del vehículo por id",
    tags=["Vehiculo"],
    description="Método para actualizar el estado (activo/inactivo) del vehículo",
    response_model=ResponseModel,
)
def updateVehiculoPorId(payload: ModificarVehiculo):
    # Preparamos la consulta para actualizar solo el campo 'Activo'
    query = f"EXEC ModificarVehiculosPorid @ID = {payload.ID}, @Activo = {payload.Activo}"

    with engine.begin() as conn:
        conn.execution_options(autocommit = True)
        roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El vehículo se actualizó de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.put(
    path="/api/vehiculo",
    name="Actualizar vehiculo",
    tags=["Vehiculo"],
    description="Método para actualizar los datos de los vehículos del cliente",
    response_model=ResponseModel,
)
def updateVehiculo(payload: VehiculoV2):
    try:
        # Combinar fotos y videos en un solo paso para optimizar el código
        campos_multimedia = [
            "MotorVehiculo_foto", "Acumulador_foto", "Espejo_retrovisor_foto", "Espejo_izquierdo_foto",
            "Espejo_derecho_foto", "Antena_foto", "Tapones_ruedas_foto", "Radio_foto", "Encendedor_foto",
            "Gato_foto", "Herramienta_foto", "Llanta_refaccion_foto", "Limpiadores_foto", "Pintura_rayada_foto",
            "Cristales_rotos_foto", "Golpes_foto", "Tapetes_foto", "Extintor_foto", "Tapones_gasolina_foto",
            "Calaveras_rotas_foto", "Molduras_completas_foto",
            "Panel_instrumentos_foto",
            "Lado_izquierdo_foto",
            "Lado_izquierdo_inf_foto",
            "Lado_derecho_foto",
            "Lado_derecho_inf_foto",
            "Tablero_foto",
            "Guantera_foto",
            "Consola_foto",
            "LadoFrontal_foto",
            "LadoTrasero_foto",
            "Cajuela_foto",
            "NumeroEconomico_foto",



            "MotorVehiculo_video", "Acumulador_video",
            "Espejo_retrovisor_video", "Espejo_izquierdo_video", "Espejo_derecho_video", "Antena_video",
            "Tapones_ruedas_video", "Radio_video", "Encendedor_video", "Gato_video", "Herramienta_video",
            "Llanta_refaccion_video", "Limpiadores_video", "Pintura_rayada_video", "Cristales_rotos_video",
            "Golpes_video", "Tapetes_video", "Extintor_video", "Tapones_gasolina_video", "Calaveras_rotas_video",
            "Molduras_completas_video","Lado_izquierdo_foto",
            "Panel_instrumentos_video",
            "Lado_izquierdo_inf_video",
            "Lado_derecho_video",
            "Lado_derecho_inf_video",
            "Tablero_video",
            "Guantera_video",
            "Consola_video",
            "LadoFrontal_video",
            "LadoTrasero_video",
            "Cajuela_video",
            "NumeroEconomico_video"
        ]

        # Crear un diccionario con los campos multimedia combinados
        multimedia = {
            campo: ",".join(getattr(payload, campo)) for campo in campos_multimedia
        }

        # Excluir los campos multimedia del payload original y combinarlos
        parametros = payload.dict(exclude=set(campos_multimedia))
        parametros.update(multimedia)

        # Consulta SQL parametrizada
        query = text("""
        exec dbo.ModificarVehiculo
            @ID = :ID,
            @Id_Cliente = :Id_Cliente,
            @Id_Empleado = :Id_Empleado,
            @Marca = :Marca,
            @Modelo = :Modelo,
            @Color = :Color,
            @No_serie = :No_serie,
            @Placa = :Placa,
            @Tipo = :Tipo,
            @Motor = :Motor,
            @Kms = :Kms,
            @MotorVehiculo = :MotorVehiculo,
            @Acumulador = :Acumulador,
            @Espejo_retrovisor = :Espejo_retrovisor,
            @Espejo_izquierdo = :Espejo_izquierdo,
            @Espejo_derecho = :Espejo_derecho,
            @Antena = :Antena,
            @Tapones_ruedas = :Tapones_ruedas,
            @Radio = :Radio,
            @Encendedor = :Encendedor,
            @Gato = :Gato,
            @Herramienta = :Herramienta,
            @Llanta_refaccion = :Llanta_refaccion,
            @Limpiadores = :Limpiadores,
            @Pintura_rayada = :Pintura_rayada,
            @Cristales_rotos = :Cristales_rotos,
            @Golpes = :Golpes,
            @Tapetes = :Tapetes,
            @Extintor = :Extintor,
            @Tapones_gasolina = :Tapones_gasolina,
            @Calaveras_rotas = :Calaveras_rotas,
            @Molduras_completas = :Molduras_completas,
            @Panel_instrumentos = :Panel_instrumentos,
            @Lado_izquierdo = :Lado_izquierdo,
            @Lado_izquierdo_inf = :Lado_izquierdo_inf,
            @Lado_derecho = :Lado_derecho,
            @Lado_derecho_inf = :Lado_derecho_inf,
            @Tablero = :Tablero,
            @Guantera = :Guantera,
            @Consola = :Consola,
            @LadoFrontal = :LadoFrontal,
            @LadoTrasero = :LadoTrasero,
            @Cajuela = :Cajuela,
            @NumeroEconomico = :NumeroEconomico,
            @MotorVehiculo_foto = :MotorVehiculo_foto,
            @Acumulador_foto = :Acumulador_foto,
            @Espejo_retrovisor_foto = :Espejo_retrovisor_foto,
            @Espejo_izquierdo_foto = :Espejo_izquierdo_foto,
            @Espejo_derecho_foto = :Espejo_derecho_foto,
            @Antena_foto = :Antena_foto,
            @Tapones_ruedas_foto = :Tapones_ruedas_foto,
            @Radio_foto = :Radio_foto,
            @Encendedor_foto = :Encendedor_foto,
            @Gato_foto = :Gato_foto,
            @Herramienta_foto = :Herramienta_foto,
            @Llanta_refaccion_foto = :Llanta_refaccion_foto,
            @Limpiadores_foto = :Limpiadores_foto,
            @Pintura_rayada_foto = :Pintura_rayada_foto,
            @Cristales_rotos_foto = :Cristales_rotos_foto,
            @Golpes_foto = :Golpes_foto,
            @Tapetes_foto = :Tapetes_foto,
            @Extintor_foto = :Extintor_foto,
            @Tapones_gasolina_foto = :Tapones_gasolina_foto,
            @Calaveras_rotas_foto = :Calaveras_rotas_foto,
            @Molduras_completas_foto = :Molduras_completas_foto,
            @Panel_instrumentos_foto = Panel_instrumentos_foto,
            @Lado_izquierdo_foto = Lado_izquierdo_foto,
            @Lado_izquierdo_inf_foto = Lado_izquierdo_inf_foto,
            @Lado_derecho_foto = Lado_derecho_foto,
            @Lado_derecho_inf_foto = Lado_derecho_inf_foto,
            @Tablero_foto = Tablero_foto,
            @Guantera_foto = Guantera_foto,
            @Consola_foto = Consola_foto,
            @LadoFrontal_foto = LadoFrontal_foto,
            @LadoTrasero_foto = LadoTrasero_foto,
            @Cajuela_foto = Cajuela_foto,
            @NumeroEconomico_foto = NumeroEconomico_foto,
            @MotorVehiculo_video = :MotorVehiculo_video,
            @Acumulador_video = :Acumulador_video,
            @Espejo_retrovisor_video = :Espejo_retrovisor_video,
            @Espejo_izquierdo_video = :Espejo_izquierdo_video,
            @Espejo_derecho_video = :Espejo_derecho_video,
            @Antena_video = :Antena_video,
            @Tapones_ruedas_video = :Tapones_ruedas_video,
            @Radio_video = :Radio_video,
            @Encendedor_video = :Encendedor_video,
            @Gato_video = :Gato_video,
            @Herramienta_video = :Herramienta_video,
            @Llanta_refaccion_video = :Llanta_refaccion_video,
            @Limpiadores_video = :Limpiadores_video,
            @Pintura_rayada_video = :Pintura_rayada_video,
            @Cristales_rotos_video = :Cristales_rotos_video,
            @Golpes_video = :Golpes_video,
            @Tapetes_video = :Tapetes_video,
            @Extintor_video = :Extintor_video,
            @Tapones_gasolina_video = :Tapones_gasolina_video,
            @Calaveras_rotas_video = :Calaveras_rotas_video,
            @Molduras_completas_video = :Molduras_completas_video,
            @Panel_instrumentos_video = Panel_instrumentos_video,
            @Lado_izquierdo_video = Lado_izquierdo_video,
            @Lado_izquierdo_inf_video = Lado_izquierdo_inf_video,
            @Lado_derecho_video = Lado_derecho_video,
            @Lado_derecho_inf_video = Lado_derecho_inf_video,
            @Tablero_video = Tablero_video,
            @Guantera_video = Guantera_video,
            @Consola_video = Consola_video,
            @LadoFrontal_video = LadoFrontal_video,
            @LadoTrasero_video = LadoTrasero_video,
            @Cajuela_video = Cajuela_video,
            @NumeroEconomico_video = NumeroEconomico_video,
            @IdFlotilla = :IdFlotilla,
            @IdOrdenServicio = :IdOrdenServicio,
            @Activo = :Activo
        """)

        # Ejecutar la consulta con los parámetros
        with engine.begin() as conn:
            result = conn.execute(query, parametros)

        # Verificar el resultado y retornar la respuesta
        if result.rowcount > 0:
            return {"id_resultado": 1, "respuesta": "Se modificó la información del vehículo de manera correcta"}
        else:
            return {"id_resultado": 0, "respuesta": "No se encontró el vehículo para modificar."}

    except Exception as e:
        # Manejo de errores
        raise HTTPException(status_code=500, detail=f"Error al modificar el vehículo: {str(e)}")


@app.put(
        path="/api/cliente",
        name='Actualizar cliente',
        tags=['Cliente'],
        description='Método para actualizar la informacion del cliente}',
        response_model=ResponseModel
)
def putcliente(payload: GetCliente ):
    query = f"exec [Clientes].[clienteinsupdel] @Accion = 3, @idCliente = {payload.ID}, @Nombre = '{payload.Nombre}', @Calle = '{payload.Calle}' \
        ,@Colonia = '{payload.Colonia}', @Ciudad = '{payload.Ciudad}',  @Estado = '{payload.Estado}', @Tel = '{payload.Tel}', @Cel = '{payload.Cel}' \
        ,@Email = '{payload.Email}', @RFC = '{payload.RFC}', @No_int = '{payload.No_int}',@Facturar_a = '{payload.Facturar_a}' \
        ,@IdUsuarioEmpleado = '{payload.Id_empleado}'"
    with engine.begin() as conn:
        conn.execution_options(autocommit = True)
        roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El cliente se actualizo de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.post(
    path="/api/ordenserviciofull",
        name='Guardar orden}',
        tags=['Orden'],
        description='Método para guardan la orden}',
        response_model=ResponseModel
)
def saveorden(payload: OrdenCompleta):

    dumpp = ResponseModel(id_resultado=1,respuesta="Se guardo la orden")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.get(
        path="/api/roles",
        name='Obtener roles',
        tags=['Catalogos'],
        description='Método para obtener los roles}',
        response_model=Roles
)
def obtener_roles():
    query = "SELECT * FROM Catalogos.Roles"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/estatus",
        name='Obtener estatus',
        tags=['Catalogos'],
        description='Método para obtener los estatus}',
    response_model=Estatus
)
def obtener_roles():
    query = "SELECT * FROM Catalogos.Estatus"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.post(
    path="/api/usuarios",
    name='crear usuario',
    tags=['Seguridad'],
    description='Método para crear usuarios',
    response_model=ResponseModel
)
async def crearusuario(payload: SaveUsuario):
    passhash = await utilsclass.get_password_hash(payload.Password)
    query = f"exec Seguridad.usuariosinsupdel @Nombre = '{payload.Nombre}', @Password = '{passhash}', @Rol = {payload.Rol}, @Estatus = {payload.Estatus}, @Accion = 1"
    with engine.begin() as conn:
        conn.execution_options(autocommit = True)
        roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="Se agregó el usuario de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.post(
    path="/api/ordenserviciopdf",
    name='obtener pdf de la orden de servicio',
    tags=['Orden'],
    description='Método para obtener el pdf de la orden de servicio',
    response_model=ResponseModel
)
def convert_html_to_pdf(clienteId: int):
    try:
        query = f"exec [Clientes].[ordendeservicio]  @idCliente = {clienteId}"
        with engine.begin() as conn:
            conn.execution_options(autocommit=True)
            roles_df = pd.read_sql(query, conn)

        # Asegúrate de que hay al menos dos filas en el DataFrame

            # Accede al valor en la segunda fila de la columna 'Nombre'


            # Construye el HTML con el valor obtenido
        orden = roles_df['idOrden'].iloc[0]
        nombre = roles_df['Nombre'].iloc[0]
        calle =roles_df['Calle'].iloc[0]
        factura =roles_df['Facturar_a'].iloc[0]
        noint = roles_df['No_int'].iloc[0]
        colonia = roles_df['Colonia'].iloc[0]
        ciuadad = roles_df['Ciudad'].iloc[0]
        estado = roles_df['Estado'].iloc[0]
        tel = roles_df['Tel'].iloc[0]
        cel = roles_df['Cel'].iloc[0]
        email = roles_df['Email'].iloc[0]
        rfc = roles_df['RFC'].iloc[0]
        marca = roles_df['Marca'].iloc[0]
        tipo = roles_df['Tipo'].iloc[0]
        modelo = roles_df['Modelo'].iloc[0]
        motor = roles_df['Motor'].iloc[0]
        color = roles_df['Color'].iloc[0]
        kms = roles_df['kms'].iloc[0]
        noserie = roles_df['No_Serie'].iloc[0]
        placa = roles_df['Placa'].iloc[0]
        espejoretrovisor = roles_df['Espejo_retrovisor'].iloc[0]
        espejoizq = roles_df['Espejo_izquierdo'].iloc[0]
        antena = roles_df['Antena'].iloc[0]
        taponesruedas = roles_df['Tapones_ruedas'].iloc[0]
        radio = roles_df['Radio'].iloc[0]
        encendedor = roles_df['Encendedor'].iloc[0]
        gato = roles_df['Gato'].iloc[0]
        herramienta = roles_df['Herramienta'].iloc[0]
        llanarefaccion = roles_df['Llanta_refaccion'].iloc[0]
        limpiadores = roles_df['Limpiadores'].iloc[0]
        pintura = roles_df['Pintura_rayada'].iloc[0]
        cristales = roles_df['Cristales_rotos'].iloc[0]
        golpes = roles_df['Golpes'].iloc[0]
        tapetes = roles_df['Tapetes'].iloc[0]
        extintor = roles_df['Extintor'].iloc[0]
        tapones_gasolina = roles_df['Tapones_gasolina'].iloc[0]
        calaveras = roles_df['Calaveras_rotas'].iloc[0]
        molduras = roles_df['Molduras_completas'].iloc[0]



        htmlstring = r"""<!DOCTYPE html>
        <html>
        <head>
        <meta charset="UTF-8">
        <div class="container">

        <div class="center-text">
        <strong> <p>ORDEN DE SERVICIO<br/>
        Refaccionaría y Taller Aautomotriz <br/>
        Trotamundos <br/>
        Horario <br/>
        Lunes a Viernes <br/>
        8 am - 6pm <br/>
        Sábados<br/>
        8 am - 3 pm<br/>
        29 Guerrero y Bravo #422.
        Col. Héroes de Nacozari. C.P. 87030 <br/>
        Tel:(834) 2285 2869 y (834) 285 2872 R.F.C. CALN810623UF3</p>
        </strong>
        </div>

        </div>

                <!-- <img style="width:30%;" align:left; src="C:\Users\arael\Downloads\trot.jpg"/>
                <img style="width:30%;" align:right; src="C:\Users\arael\Downloads\todo.jpg"/> -->
                <!-- <h2 style="text-align:center;font-size:25px;">ORDEN DE SERVICIO</h2>
                <h3 style="text-align:center;">Refaccionaría y Taller Aautomotriz</h3>
                <div style="text-align: right;">Horario <br/></h3>
                Lunes a Viernes <br/>
                8 am - 6pm <br/>
                Sábados<br/>
                8 am - 3 pm<br/>
                <h3 style="text-align:center;">Trotamundos</h3>
                <h3 style="text-align:center;">29 Guerrero y Bravo #422.</h3>
                <h3 style="text-align:center;">Col. Héroes de Nacozari. C.P. 87030</h3>
                <h3 style="text-align:center;">Tel:(834) 2285 2869 y (834) 285 2872 R.F.C. CALN810623UF3</h3> -->
                <div style="position: relative;">
                <div style="width: 20%; position: relative;">"""
        htmlstring +=f"""
                <table>
                <th>
                  ORDEN
                </th>
                <tr>
                <td>
        {roles_df['idOrden'].iloc[0]}
                </td>
                </tr>
                <tr>
                <td>
        Fecha de Recepción:
            </td>
                </tr>
                <tr>
                <td>
        Fecha de Entrega:
                </td>
                </tr>
                </table>
                </div>
                <div style="width: 100%; position: absolute; left: 30%; top: 0; height: 100%;">
                <table>
                <tr>
        <td>Hora de compromiso</td>
        <td>Motivo de Visita:  (Previsto)  (Correctivo)</td>
                </tr>
                <tr>
        <td>Hora de entrega</td>
        <td>Medio:  Periódico  Radio  TV  Volante   Recomendación  Otros</td>
                </tr>
                </table>
                </div>
                </div>
                </head>
                <meta charset="UTF-8">"""
        htmlstring +="""
                <style>
                .left-image, .right-image {
        max-width: 20%;
        height: 35%;
        border-radius:50%;
        }
                .container {
        display: flex;
        justify-content: space-between;
        align-items: center;
        width: 100%;
        }
        .left-image, .right-image {
        max-width: 20%; /* Puedes ajustar el tamaño de las imágenes según tus necesidades */
        height: auto;
        }
        .center-text {
        flex: 1;
        text-align: center;
        font-size: 20px;
        padding: 0 20px; /* Espaciado alrededor del texto */
        font-weight: 50%;
        }
                table {
        font-family: arial, sans-serif;
        border-collapse: collapse;
        width: 100%;
        border: 1px solid #000;
                }
                td, th {
        border-collapse: collapse; border: 1px solid #dddddd;
        text-align: left;
        border: 1px solid #000;
        vertical-align: top;
        padding: 8px;
                }
                th {
        background-color: #dddddd;
                }
                .linea {
                border-top: 1px solid black;
                height: 2px;
                max-width: 200px;
                padding: 0;
                margin: 20px auto 0 auto;
                }
                td.empty-cell {
        min-width: 50px; /* Establece el tamaño mínimo para celdas vacías */
        min-height: 20px; /* Establece la altura mínima para celdas vacías */
                }
                </style>"""
        htmlstring += f"""
                <body>
                <div style="position: relative;">
        <div style="width: 40%; position: relative; height: 100%;">
        <table>
        <th>Cliente</th>
        <tr>
        <td>Facturar a:{factura}</td>
        </tr>
        <tr>
        <td>Nombre:{nombre}</td>
        </tr>
        <tr>
        <td>Calle:{calle} </td>
        <td>No. int:{noint}</td>
        </tr>
        <tr>
        <td>Colonia:{colonia}</td>
        <td>Ciudad:{ciuadad}</td>
        </tr>
        <tr>
        <td>Cumpleaños:</td>
        <td>Estado:{estado}</td>
        </tr>
        <tr>
        <td>Tel:{tel}</td>
        <td>Cel. /Nex:{cel} </td>
        </tr>
        <tr>
        <td>EMAIL:{email}</td>
        </tr>
        <tr>
        <td>RFC:{rfc}</td>
        </tr>



        </table>
        </div>
        <div style="width: 50%; position: absolute; left: 50%; top: 0; height: 100%;">
        <table>
        <th>Vehículo</th>
        <tr>
        <td>Marca:{marca}</td>
        <td>Tipo:{tipo}</td>
        </tr>
        <tr>
        <td>Modelo:{modelo}</td>
        <td>Motor:{motor}</td>
        </tr>
        <tr>
        <td>Color:{color}</td>
        <td>Kms:{kms}</td>
        </tr>
        <tr>
        <td>N. Serie:{noserie}</td>
        </tr>
        <tr>
          <td>Placa:{placa}</td>
        </tr>
        <tr>
          <th>Inventario de Vehículo</th>
        </tr>
        <tr>
        <td>Espejo Retrovisor:</td>
        <td>{espejoretrovisor}</td>
        </tr>
        <tr>
        <td>Espejo Izquierdo:</td>
        <td>{espejoizq}</td>
        </tr>
        <tr>
        <td>Espejo Derecho:</td>
        <td></td>
        </tr>
        <tr>
        <td>Antena:</td>
        <td>{antena}</td>
        </tr>
        <tr>
        <td>Tapones de Ruedas:</td>
        <td>{taponesruedas}</td>
        </tr>
        <tr>
        <tr>
        <td>Radio:</td>
        <td>{radio}</td>
        </tr>
        <tr>
        <td>Encendedor:</td>
        <td>{encendedor}</td>
        </tr>
        <tr>
        <td>Gato:</td>
        <td>{gato}</td>
        </tr>
        <tr>
        <td>Herramienta:</td>
        <td>{herramienta}</td>
        </tr>
        <tr>
        <td>Llanta de Refacción:</td>
        <td>{llanarefaccion}</td>
        </tr>
        <tr>
        <td>Limpiadores:</td>
        <td>{limpiadores}</td>
        </tr>
        <tr>
        <td>Pintura Rayada:</td>
        <td>{pintura}</td>
        </tr>
        <tr>
        <td>Cristales Rotos:</td>
        <td>{cristales}</td>
        </tr>
        <tr>
        <td>Golpes:</td>
        <td>{golpes}</td>
        </tr>
        <tr>
        <td>Tapetes:</td>
        <td>{tapetes}</td>
        </tr>
        <tr>
        <td>Extintor:</td>
        <td>{extintor}</td>
        </tr>
        <tr>
        <td>Tapón de Gasolina:</td>
        <td>{tapones_gasolina}</td>
        </tr>
        <tr>
        <td>Calaveras Rotas:</td>
        <td>{calaveras}</td>
        </tr>
        <tr>
        <td>Molduras Completas:</td>
        <td>{molduras}</td>
        </tr>
                <div></div>
                <table>
                <tr>
                <td>
                Firma del Proveedor
                </td>
                <td>
                Firma de Aceptación del Cliente
                </td>
                </tr>
                </table>
                <div></div>
                <table>
                <th>
                Servicio Solicitado
                </th>
                <th>
                Recibió
                </th>
                <th>
                Técnico
                </th>
                <th>
                Orden
                </th>
                <tr>
                  <td></td>
                </tr>
                </table>
                </body>
                </html>"""
         # Resto del código para convertir a PDF...
        img = "\\img1.jpg"
        pdf_path = "example.pdf"
          # Rutas y configuraciones para Linux
        path_wkhtmltopdf = '/usr/local/bin/wkhtmltopdf'
        config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)
        pdfkit.from_string(htmlstring, 'reporte.pdf', configuration=config)
        return JSONResponse(content={"message": "PDF creado exitosamente"}, status_code=200)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get(
        path="/api/empleados",
        name='Obtener empleados',
        tags=['Seguridad'],
        description='Método para obtener la informacion de todos los empleados',
        response_model=List[Empleado]
)
def getempleados():
    query = f"[Seguridad].[usuariosinsupdel] @Accion = 2"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/empleado",
        name='Obtener empleado',
        tags=['Seguridad'],
        description='Método para obtener la informacion de todos los empleados',
        response_model=Empleado
)
def getempleados(IdUsuario: str):
    query = f"[Seguridad].[usuariosinsupdel] @Accion = 3, @Idusuario = {IdUsuario}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.post(
        path="/api/checklisthistorico",
        name='Insertar checklist histórico',
        tags=['ChecklistHistorico'],
        description='Método para insertar el histórico del checklist',
        response_model=CheckListHistorico
)
def savechecklisthistorico(payload: CheckListHistorico):
    query = f"""
    EXEC InsertarHistoricoCheckList
        @IdChecklist = {payload.IdChecklist}, \
        @IdVehiculo = {payload.IdVehiculo}, \
        @IdEmpleado = {payload.IdEmpleado}, \
        @Fecha = '{payload.Fecha}', \
        @TiempoTranscurrido = {payload.TiempoTranscurrido}, \
        @Estado = '{payload.Estado}', \
       """
    print (query)
    with engine.begin() as conn:
          conn.execution_options(autocommit = True)
          roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El histórico del checklist se guardó de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.post(
        path="/api/checklisthistoricoservicio",
        name='Insertar checklist histórico servicio',
        tags=['ChecklistHistorico'],
        description='Método para insertar el histórico del checklist servicio',
        response_model=CheckListHistorico
)
def savechecklisthistoricoservicio(payload: CheckListHistorico):
    query = f"""
    EXEC InsertarHistoricoCheckListServicio
        @IdChecklist = {payload.IdChecklist}, \
        @IdVehiculo = {payload.IdVehiculo}, \
        @IdEmpleado = {payload.IdEmpleado}, \
        @Fecha = '{payload.Fecha}', \
        @TiempoTranscurrido = {payload.TiempoTranscurrido}, \
        @Estado = '{payload.Estado}', \
       """
    print (query)
    with engine.begin() as conn:
          conn.execution_options(autocommit = True)
          roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El histórico del checklist se guardó de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.post(
        path="/api/checklist",
        name='Insertar checklist',
        tags=['Checklist'],
        description='Método para insertar el checklist',
        response_model=Checklist
)
def savechecklist(payload: Checklist):
    query = f"""
    EXEC InsertChecklist
        @lectura_codigos = {payload.lectura_codigos}, \
        @lectura_codigos_observacion = N'{payload.lectura_codigos_observacion}' , \
        @lectura_codigos_foto = N'{payload.lectura_codigos_foto}' , \
        @servofreno = {payload.servofreno}, \
        @servofreno_observacion = N'{payload.servofreno_observacion}' , \
        @servofreno_foto = N'{payload.servofreno_foto}' , \
        @pedal_freno = {payload.pedal_freno}, \
        @pedal_freno_observacion = N'{payload.pedal_freno_observacion}' , \
        @pedal_freno_foto = N'{payload.pedal_freno_foto}' , \
        @pedal_estacionamiento = {payload.pedal_estacionamiento}, \
        @pedal_estacionamiento_observacion = N'{payload.pedal_estacionamiento_observacion}' , \
        @pedal_estacionamiento_foto = N'{payload.pedal_estacionamiento_foto}' , \
        @cinturon_seguridad = {payload.cinturon_seguridad}, \
        @cinturon_seguridad_observacion = N'{payload.cinturon_seguridad_observacion}' , \
        @cinturon_seguridad_foto = N'{payload.cinturon_seguridad_foto}' , \
        @cuadro_instrumentos = {payload.cuadro_instrumentos}, \
        @cuadro_instrumentos_observacion = N'{payload.cuadro_instrumentos_observacion}' ,\
        @cuadro_instrumentos_foto = N'{payload.cuadro_instrumentos_foto}' ,\
        @aire_acondicionado = {payload.aire_acondicionado},\
        @aire_acondicionado_observacion = N'{payload.aire_acondicionado_observacion}' ,\
        @aire_acondicionado_foto = N'{payload.aire_acondicionado_foto}' ,\
        @bocina_claxon = {payload.bocina_claxon},\
        @bocina_claxon_observacion = N'{payload.bocina_claxon_observacion}' ,\
        @bocina_claxon_foto = N'{payload.bocina_claxon_foto}' ,\
        @iluminacion_interior = {payload.iluminacion_interior},\
        @iluminacion_interior_observacion = N'{payload.iluminacion_interior_observacion}' ,\
        @iluminacion_interior_foto = N'{payload.iluminacion_interior_foto}' ,\
        @iluminacion_externa = {payload.iluminacion_externa},\
        @iluminacion_externa_observacion = N'{payload.iluminacion_externa_observacion}' ,\
        @iluminacion_externa_foto = N'{payload.iluminacion_externa_foto}' ,\
        @limpiaparabrisas = {payload.limpiaparabrisas}, \
        @limpiaparabrisas_observacion = N'{payload.limpiaparabrisas_observacion}' ,\
        @limpiaparabrisas_foto = N'{payload.limpiaparabrisas_foto}' , \
        @limpia_medallon = {payload.limpia_medallon}, \
        @limpia_medallon_observacion = N'{payload.limpia_medallon_observacion}' , \
        @limpia_medallon_foto = N'{payload.limpia_medallon_foto}' , \
        @neumaticos_friccion = {payload.neumaticos_friccion}, \
        @neumaticos_friccion_observacion = N'{payload.neumaticos_friccion_observacion}' ,  \
        @neumaticos_friccion_foto = N'{payload.neumaticos_friccion_foto}' , \
        @otros_vehiculo_en_piso = {payload.otros_vehiculo_en_piso}, \
        @otros_vehiculo_en_piso_observacion = N'{payload.otros_vehiculo_en_piso_observacion}' , \
        @otros_vehiculo_en_piso_foto = N'{payload.otros_vehiculo_en_piso_foto}' , \
        @estado_fugas_aceite = {payload.estado_fugas_aceite}, \
        @estado_fugas_aceite_observacion = N'{payload.estado_fugas_aceite_observacion}' , \
        @estado_fugas_aceite_foto = N'{payload.estado_fugas_aceite_foto}' , \
        @estado_nivel_calidad_lubricante_transmision = {payload.estado_nivel_calidad_lubricante_transmision}, \
        @estado_nivel_calidad_lubricante_transmision_observacion = N'{payload.estado_nivel_calidad_lubricante_transmision_observacion}' , \
        @estado_nivel_calidad_lubricante_transmision_foto = N'{payload.estado_nivel_calidad_lubricante_transmision_foto}' , \
        @estado_nivel_calidad_lubricante_diferencial = {payload.estado_nivel_calidad_lubricante_diferencial}, \
        @estado_nivel_calidad_lubricante_diferencial_observacion = N'{payload.estado_nivel_calidad_lubricante_diferencial_observacion}' , \
        @estado_nivel_calidad_lubricante_diferencial_foto = N'{payload.estado_nivel_calidad_lubricante_diferencial_foto}' , \
        @cubrepolvos_flechas = {payload.cubrepolvos_flechas}, \
        @cubrepolvos_flechas_observacion = N'{payload.cubrepolvos_flechas_observacion}' , \
        @cubrepolvos_flechas_foto = N'{payload.cubrepolvos_flechas_foto}' , \
        @componentes_direccion = {payload.componentes_direccion}, \
        @componentes_direccion_observacion = N'{payload.componentes_direccion_observacion}' , \
        @componentes_direccion_foto = N'{payload.componentes_direccion_foto}' , \
        @componentes_suspesion = {payload.componentes_suspesion}, \
        @componentes_suspesion_observacion = N'{payload.componentes_suspesion_observacion}' , \
        @componentes_suspesion_foto = N'{payload.componentes_suspesion_foto}' , \
        @sistema_escape_completo = {payload.sistema_escape_completo}, \
        @sistema_escape_completo_observacion = N'{payload.sistema_escape_completo_observacion}' , \
        @sistema_escape_completo_foto = N'{payload.sistema_escape_completo_foto}' , \
        @sistema_alimentacion_combustible = {payload.sistema_alimentacion_combustible}, \
        @sistema_alimentacion_combustible_observacion = N'{payload.sistema_alimentacion_combustible_observacion}' , \
        @sistema_alimentacion_combustible_foto = N'{payload.sistema_alimentacion_combustible_foto}' , \
        @filtro_combustible = {payload.filtro_combustible}, \
        @filtro_combustible_observacion = N'{payload.filtro_combustible_observacion}' , \
        @filtro_combustible_foto = N'{payload.filtro_combustible_foto}' , \
        @control_fugas_direccion_hidraulica = {payload.control_fugas_direccion_hidraulica}, \
        @control_fugas_direccion_hidraulica_observacion = N'{payload.control_fugas_direccion_hidraulica_observacion}' , \
        @control_fugas_direccion_hidraulica_foto = N'{payload.control_fugas_direccion_hidraulica_foto}' , \
        @otros_altura_total = {payload.otros_altura_total}, \
        @otros_altura_total_observacion = N'{payload.otros_altura_total_observacion}' , \
        @otros_altura_total_foto = N'{payload.otros_altura_total_foto}' , \
        @rodamiento_mazas_rueda = {payload.rodamiento_mazas_rueda}, \
        @rodamiento_mazas_rueda_observacion = N'{payload.rodamiento_mazas_rueda_observacion}' , \
        @rodamiento_mazas_rueda_foto = N'{payload.rodamiento_mazas_rueda_foto}' , \
        @holgura_partes_suspension_rueda = {payload.holgura_partes_suspension_rueda}, \
        @holgura_partes_suspension_rueda_observacion = N'{payload.holgura_partes_suspension_rueda_observacion}' , \
        @holgura_partes_suspension_rueda_foto = N'{payload.holgura_partes_suspension_rueda_foto}' , \
        @control_neumaticos_desgaste_presion = {payload.control_neumaticos_desgaste_presion}, \
        @control_neumaticos_desgaste_presion_observacion = N'{payload.control_neumaticos_desgaste_presion_observacion}' , \
        @control_neumaticos_desgaste_presion_foto = N'{payload.control_neumaticos_desgaste_presion_foto}' , \
        @profundidad = {payload.profundidad}, \
        @profundidad_observacion = N'{payload.profundidad_observacion}' , \
        @profundidad_foto = N'{payload.profundidad_foto}' , \
        @presion = {payload.presion}, \
        @presion_observacion = N'{payload.presion_observacion}' , \
        @presion_foto = N'{payload.presion_foto}' , \
        @otros_altura_media = {payload.otros_altura_media}, \
        @otros_altura_media_observacion = N'{payload.otros_altura_media_observacion}' , \
        @otros_altura_media_foto = N'{payload.otros_altura_media_foto}' , \
        @nivel_calidad_aceite_motor = {payload.nivel_calidad_aceite_motor}, \
        @nivel_calidad_aceite_motor_observacion = N'{payload.nivel_calidad_aceite_motor_observacion}' , \
        @nivel_calidad_aceite_motor_foto = N'{payload.nivel_calidad_aceite_motor_foto}' , \
        @filtro_aire = {payload.filtro_aire}, \
        @filtro_aire_observacion = N'{payload.filtro_aire_observacion}' , \
        @filtro_aire_foto = N'{payload.filtro_aire_foto}' , \
        @filtro_polen = {payload.filtro_polen}, \
        @filtro_polen_observacion = N'{payload.filtro_polen_observacion}' , \
        @filtro_polen_foto = N'{payload.filtro_polen_foto}' , \
        @filtro_pcv = {payload.filtro_pcv}, \
        @filtro_pcv_observacion = N'{payload.filtro_pcv_observacion}' , \
        @filtro_pcv_foto = N'{payload.filtro_pcv_foto}' , \
        @valvula_pcv = {payload.valvula_pcv}, \
        @valvula_pcv_observacion = N'{payload.valvula_pcv_observacion}' , \
        @valvula_pcv_foto = N'{payload.valvula_pcv_foto}' , \
        @bujias_encendido = {payload.bujias_encendido}, \
        @bujias_encendido_observacion = N'{payload.bujias_encendido_observacion}' , \
        @bujias_encendido_foto = N'{payload.bujias_encendido_foto}' , \
        @cables_bujias_bobinas_ignicion = {payload.cables_bujias_bobinas_ignicion}, \
        @cables_bujias_bobinas_ignicion_observacion = N'{payload.cables_bujias_bobinas_ignicion_observacion}' , \
        @cables_bujias_bobinas_ignicion_foto = N'{payload.cables_bujias_bobinas_ignicion_foto}' , \
        @nivel_anticongenlante = {payload.nivel_anticongenlante}, \
        @nivel_anticongenlante_observacion = N'{payload.nivel_anticongenlante_observacion}' , \
        @nivel_anticongenlante_foto = N'{payload.nivel_anticongenlante_foto}' , \
        @tapon_radiador = {payload.tapon_radiador}, \
        @tapon_radiador_observacion = N'{payload.tapon_radiador_observacion}' , \
        @tapon_radiador_foto = N'{payload.tapon_radiador_foto}' , \
        @mangueras_sistema = {payload.mangueras_sistema}, \
        @mangueras_sistema_observacion = N'{payload.mangueras_sistema_observacion}' , \
        @mangueras_sistema_foto = N'{payload.mangueras_sistema_foto}' , \
        @desempeño_ventilador = {payload.desempeño_ventilador}, \
        @desempeño_ventilador_observacion = N'{payload.desempeño_ventilador_observacion}' , \
        @desempeño_ventilador_foto = N'{payload.desempeño_ventilador_foto}' , \
        @calidad_liquido_limpiaparabrisas = {payload.calidad_liquido_limpiaparabrisas}, \
        @calidad_liquido_limpiaparabrisas_observacion = N'{payload.calidad_liquido_limpiaparabrisas_observacion}' , \
        @calidad_liquido_limpiaparabrisas_foto = N'{payload.calidad_liquido_limpiaparabrisas_foto}' , \
        @calidad_aceite_direccion_hidraulica = {payload.calidad_aceite_direccion_hidraulica}, \
        @calidad_aceite_direccion_hidraulica_observacion = N'{payload.calidad_aceite_direccion_hidraulica_observacion}' , \
        @calidad_aceite_direccion_hidraulica_foto = N'{payload.calidad_aceite_direccion_hidraulica_foto}' , \
        @calidad_aceite_transmision_bayoneta = {payload.calidad_aceite_transmision_bayoneta}, \
        @calidad_aceite_transmision_bayoneta_observacion = N'{payload.calidad_aceite_transmision_bayoneta_observacion}' , \
        @calidad_aceite_transmision_bayoneta_foto = N'{payload.calidad_aceite_transmision_bayoneta_foto}' , \
        @liquido_bateria_condiciones = {payload.liquido_bateria_condiciones}, \
        @liquido_bateria_condiciones_observacion = N'{payload.liquido_bateria_condiciones_observacion}' , \
        @liquido_bateria_condiciones_foto = N'{payload.liquido_bateria_condiciones_foto}' ,    \
        @bandas_poly_v = {payload.bandas_poly_v}, \
        @bandas_poly_v_observacion = N'{payload.bandas_poly_v_observacion}' , \
        @bandas_poly_v_foto = N'{payload.bandas_poly_v_foto}' , \
        @poleas_banda = {payload.poleas_banda}, \
        @poleas_banda_observacion = N'{payload.poleas_banda_observacion}' , \
        @poleas_banda_foto = N'{payload.poleas_banda_foto}' , \
        @banda_tiempo = {payload.banda_tiempo}, \
        @banda_tiempo_observacion = N'{payload.banda_tiempo_observacion}' , \
        @banda_tiempo_foto = N'{payload.banda_tiempo_foto}' , \
        @otros_habitaculo_motor = {payload.otros_habitaculo_motor}, \
        @otros_habitaculo_motor_observacion = N'{payload.otros_habitaculo_motor_observacion}' , \
        @otros_habitaculo_motor_foto = N'{payload.otros_habitaculo_motor_foto}' , \
        @reset_intervalo_servicio = {payload.reset_intervalo_servicio}, \
        @reset_intervalo_servicio_observacion = N'{payload.reset_intervalo_servicio_observacion}' , \
        @reset_intervalo_servicio_foto = N'{payload.reset_intervalo_servicio_foto}' , \
        @ajuste_tornillos_neumaticos_torquimetro = {payload.ajuste_tornillos_neumaticos_torquimetro}, \
        @ajuste_tornillos_neumaticos_torquimetro_observacion = N'{payload.ajuste_tornillos_neumaticos_torquimetro_observacion}' , \
        @ajuste_tornillos_neumaticos_torquimetro_foto = N'{payload.ajuste_tornillos_neumaticos_torquimetro_foto}' , \
        @limpiar_libricar_puertas_cerraduras = {payload.limpiar_libricar_puertas_cerraduras}, \
        @limpiar_libricar_puertas_cerraduras_observacion = N'{payload.limpiar_libricar_puertas_cerraduras_observacion}' , \
        @limpiar_libricar_puertas_cerraduras_foto = N'{payload.limpiar_libricar_puertas_cerraduras_foto}' , \
        @completar_plan_mantenimiento = {payload.completar_plan_mantenimiento}, \
        @completar_plan_mantenimiento_observacion = N'{payload.completar_plan_mantenimiento_observacion}' , \
        @completar_plan_mantenimiento_foto = N'{payload.completar_plan_mantenimiento_foto}' , \
        @fecha = N'{payload.Fecha}' , \
        @Id_empleado = {payload.IdEmpleado} , \
        @Id_vehiculo = {payload.IdVehiculo} , \
        @id_ordendeservicio = {payload.Id_ordendeservicio} , \
        @NumeroSerie = '{payload.NumeroSerie}' """
    print (query)
    with engine.begin() as conn:
          conn.execution_options(autocommit = True)
          roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El checklist se guardó de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)

@app.put(
    path="/api/checklist",
    name="Actualizar checklist",
    tags=["Checklist"],
    description="Método para actualizar la información del checklist",
    response_model=Checklist
)
def put_checklist(payload: Checklist):
    # Construcción de la consulta SQL para ejecutar el procedimiento almacenado
    query = f"""
    exec [dbo].[UpdateCheckList]
        @lectura_codigos = {payload.lectura_codigos}, \
        @lectura_codigos_observacion = N'{payload.lectura_codigos_observacion}' , \
        @lectura_codigos_foto = N'{payload.lectura_codigos_foto}' , \
        @servofreno = {payload.servofreno}, \
        @servofreno_observacion = N'{payload.servofreno_observacion}' , \
        @servofreno_foto = N'{payload.servofreno_foto}' , \
        @pedal_freno = {payload.pedal_freno}, \
        @pedal_freno_observacion = N'{payload.pedal_freno_observacion}' , \
        @pedal_freno_foto = N'{payload.pedal_freno_foto}' , \
        @pedal_estacionamiento = {payload.pedal_estacionamiento}, \
        @pedal_estacionamiento_observacion = N'{payload.pedal_estacionamiento_observacion}' , \
        @pedal_estacionamiento_foto = N'{payload.pedal_estacionamiento_foto}' , \
        @cinturon_seguridad = {payload.cinturon_seguridad}, \
        @cinturon_seguridad_observacion = N'{payload.cinturon_seguridad_observacion}' , \
        @cinturon_seguridad_foto = N'{payload.cinturon_seguridad_foto}' , \
        @cuadro_instrumentos = {payload.cuadro_instrumentos}, \
        @cuadro_instrumentos_observacion = N'{payload.cuadro_instrumentos_observacion}' ,\
        @cuadro_instrumentos_foto = N'{payload.cuadro_instrumentos_foto}' ,\
        @aire_acondicionado = {payload.aire_acondicionado},\
        @aire_acondicionado_observacion = N'{payload.aire_acondicionado_observacion}' ,\
        @aire_acondicionado_foto = N'{payload.aire_acondicionado_foto}' ,\
        @bocina_claxon = {payload.bocina_claxon},\
        @bocina_claxon_observacion = N'{payload.bocina_claxon_observacion}' ,\
        @bocina_claxon_foto = N'{payload.bocina_claxon_foto}' ,\
        @iluminacion_interior = {payload.iluminacion_interior},\
        @iluminacion_interior_observacion = N'{payload.iluminacion_interior_observacion}' ,\
        @iluminacion_interior_foto = N'{payload.iluminacion_interior_foto}' ,\
        @iluminacion_externa = {payload.iluminacion_externa},\
        @iluminacion_externa_observacion = N'{payload.iluminacion_externa_observacion}' ,\
        @iluminacion_externa_foto = N'{payload.iluminacion_externa_foto}' ,\
        @limpiaparabrisas = {payload.limpiaparabrisas}, \
        @limpiaparabrisas_observacion = N'{payload.limpiaparabrisas_observacion}' ,\
        @limpiaparabrisas_foto = N'{payload.limpiaparabrisas_foto}' , \
        @limpia_medallon = {payload.limpia_medallon}, \
        @limpia_medallon_observacion = N'{payload.limpia_medallon_observacion}' , \
        @limpia_medallon_foto = N'{payload.limpia_medallon_foto}' , \
        @neumaticos_friccion = {payload.neumaticos_friccion}, \
        @neumaticos_friccion_observacion = N'{payload.neumaticos_friccion_observacion}' ,  \
        @neumaticos_friccion_foto = N'{payload.neumaticos_friccion_foto}' , \
        @otros_vehiculo_en_piso = {payload.otros_vehiculo_en_piso}, \
        @otros_vehiculo_en_piso_observacion = N'{payload.otros_vehiculo_en_piso_observacion}' , \
        @otros_vehiculo_en_piso_foto = N'{payload.otros_vehiculo_en_piso_foto}' , \
        @estado_fugas_aceite = {payload.estado_fugas_aceite}, \
        @estado_fugas_aceite_observacion = N'{payload.estado_fugas_aceite_observacion}' , \
        @estado_fugas_aceite_foto = N'{payload.estado_fugas_aceite_foto}' , \
        @estado_nivel_calidad_lubricante_transmision = {payload.estado_nivel_calidad_lubricante_transmision}, \
        @estado_nivel_calidad_lubricante_transmision_observacion = N'{payload.estado_nivel_calidad_lubricante_transmision_observacion}' , \
        @estado_nivel_calidad_lubricante_transmision_foto = N'{payload.estado_nivel_calidad_lubricante_transmision_foto}' , \
        @estado_nivel_calidad_lubricante_diferencial = {payload.estado_nivel_calidad_lubricante_diferencial}, \
        @estado_nivel_calidad_lubricante_diferencial_observacion = N'{payload.estado_nivel_calidad_lubricante_diferencial_observacion}' , \
        @estado_nivel_calidad_lubricante_diferencial_foto = N'{payload.estado_nivel_calidad_lubricante_diferencial_foto}' , \
        @cubrepolvos_flechas = {payload.cubrepolvos_flechas}, \
        @cubrepolvos_flechas_observacion = N'{payload.cubrepolvos_flechas_observacion}' , \
        @cubrepolvos_flechas_foto = N'{payload.cubrepolvos_flechas_foto}' , \
        @componentes_direccion = {payload.componentes_direccion}, \
        @componentes_direccion_observacion = N'{payload.componentes_direccion_observacion}' , \
        @componentes_direccion_foto = N'{payload.componentes_direccion_foto}' , \
        @componentes_suspesion = {payload.componentes_suspesion}, \
        @componentes_suspesion_observacion = N'{payload.componentes_suspesion_observacion}' , \
        @componentes_suspesion_foto = N'{payload.componentes_suspesion_foto}' , \
        @sistema_escape_completo = {payload.sistema_escape_completo}, \
        @sistema_escape_completo_observacion = N'{payload.sistema_escape_completo_observacion}' , \
        @sistema_escape_completo_foto = N'{payload.sistema_escape_completo_foto}' , \
        @sistema_alimentacion_combustible = {payload.sistema_alimentacion_combustible}, \
        @sistema_alimentacion_combustible_observacion = N'{payload.sistema_alimentacion_combustible_observacion}' , \
        @sistema_alimentacion_combustible_foto = N'{payload.sistema_alimentacion_combustible_foto}' , \
        @filtro_combustible = {payload.filtro_combustible}, \
        @filtro_combustible_observacion = N'{payload.filtro_combustible_observacion}' , \
        @filtro_combustible_foto = N'{payload.filtro_combustible_foto}' , \
        @control_fugas_direccion_hidraulica = {payload.control_fugas_direccion_hidraulica}, \
        @control_fugas_direccion_hidraulica_observacion = N'{payload.control_fugas_direccion_hidraulica_observacion}' , \
        @control_fugas_direccion_hidraulica_foto = N'{payload.control_fugas_direccion_hidraulica_foto}' , \
        @otros_altura_total = {payload.otros_altura_total}, \
        @otros_altura_total_observacion = N'{payload.otros_altura_total_observacion}' , \
        @otros_altura_total_foto = N'{payload.otros_altura_total_foto}' , \
        @rodamiento_mazas_rueda = {payload.rodamiento_mazas_rueda}, \
        @rodamiento_mazas_rueda_observacion = N'{payload.rodamiento_mazas_rueda_observacion}' , \
        @rodamiento_mazas_rueda_foto = N'{payload.rodamiento_mazas_rueda_foto}' , \
        @holgura_partes_suspension_rueda = {payload.holgura_partes_suspension_rueda}, \
        @holgura_partes_suspension_rueda_observacion = N'{payload.holgura_partes_suspension_rueda_observacion}' , \
        @holgura_partes_suspension_rueda_foto = N'{payload.holgura_partes_suspension_rueda_foto}' , \
        @control_neumaticos_desgaste_presion = {payload.control_neumaticos_desgaste_presion}, \
        @control_neumaticos_desgaste_presion_observacion = N'{payload.control_neumaticos_desgaste_presion_observacion}' , \
        @control_neumaticos_desgaste_presion_foto = N'{payload.control_neumaticos_desgaste_presion_foto}' , \
        @profundidad = {payload.profundidad}, \
        @profundidad_observacion = N'{payload.profundidad_observacion}' , \
        @profundidad_foto = N'{payload.profundidad_foto}' , \
        @presion = {payload.presion}, \
        @presion_observacion = N'{payload.presion_observacion}' , \
        @presion_foto = N'{payload.presion_foto}' , \
        @otros_altura_media = {payload.otros_altura_media}, \
        @otros_altura_media_observacion = N'{payload.otros_altura_media_observacion}' , \
        @otros_altura_media_foto = N'{payload.otros_altura_media_foto}' , \
        @nivel_calidad_aceite_motor = {payload.nivel_calidad_aceite_motor}, \
        @nivel_calidad_aceite_motor_observacion = N'{payload.nivel_calidad_aceite_motor_observacion}' , \
        @nivel_calidad_aceite_motor_foto = N'{payload.nivel_calidad_aceite_motor_foto}' , \
        @filtro_aire = {payload.filtro_aire}, \
        @filtro_aire_observacion = N'{payload.filtro_aire_observacion}' , \
        @filtro_aire_foto = N'{payload.filtro_aire_foto}' , \
        @filtro_polen = {payload.filtro_polen}, \
        @filtro_polen_observacion = N'{payload.filtro_polen_observacion}' , \
        @filtro_polen_foto = N'{payload.filtro_polen_foto}' , \
        @filtro_pcv = {payload.filtro_pcv}, \
        @filtro_pcv_observacion = N'{payload.filtro_pcv_observacion}' , \
        @filtro_pcv_foto = N'{payload.filtro_pcv_foto}' , \
        @valvula_pcv = {payload.valvula_pcv}, \
        @valvula_pcv_observacion = N'{payload.valvula_pcv_observacion}' , \
        @valvula_pcv_foto = N'{payload.valvula_pcv_foto}' , \
        @bujias_encendido = {payload.bujias_encendido}, \
        @bujias_encendido_observacion = N'{payload.bujias_encendido_observacion}' , \
        @bujias_encendido_foto = N'{payload.bujias_encendido_foto}' , \
        @cables_bujias_bobinas_ignicion = {payload.cables_bujias_bobinas_ignicion}, \
        @cables_bujias_bobinas_ignicion_observacion = N'{payload.cables_bujias_bobinas_ignicion_observacion}' , \
        @cables_bujias_bobinas_ignicion_foto = N'{payload.cables_bujias_bobinas_ignicion_foto}' , \
        @nivel_anticongenlante = {payload.nivel_anticongenlante}, \
        @nivel_anticongenlante_observacion = N'{payload.nivel_anticongenlante_observacion}' , \
        @nivel_anticongenlante_foto = N'{payload.nivel_anticongenlante_foto}' , \
        @tapon_radiador = {payload.tapon_radiador}, \
        @tapon_radiador_observacion = N'{payload.tapon_radiador_observacion}' , \
        @tapon_radiador_foto = N'{payload.tapon_radiador_foto}' , \
        @mangueras_sistema = {payload.mangueras_sistema}, \
        @mangueras_sistema_observacion = N'{payload.mangueras_sistema_observacion}' , \
        @mangueras_sistema_foto = N'{payload.mangueras_sistema_foto}' , \
        @desempeño_ventilador = {payload.desempeño_ventilador}, \
        @desempeño_ventilador_observacion = N'{payload.desempeño_ventilador_observacion}' , \
        @desempeño_ventilador_foto = N'{payload.desempeño_ventilador_foto}' , \
        @calidad_liquido_limpiaparabrisas = {payload.calidad_liquido_limpiaparabrisas}, \
        @calidad_liquido_limpiaparabrisas_observacion = N'{payload.calidad_liquido_limpiaparabrisas_observacion}' , \
        @calidad_liquido_limpiaparabrisas_foto = N'{payload.calidad_liquido_limpiaparabrisas_foto}' , \
        @calidad_aceite_direccion_hidraulica = {payload.calidad_aceite_direccion_hidraulica}, \
        @calidad_aceite_direccion_hidraulica_observacion = N'{payload.calidad_aceite_direccion_hidraulica_observacion}' , \
        @calidad_aceite_direccion_hidraulica_foto = N'{payload.calidad_aceite_direccion_hidraulica_foto}' , \
        @calidad_aceite_transmision_bayoneta = {payload.calidad_aceite_transmision_bayoneta}, \
        @calidad_aceite_transmision_bayoneta_observacion = N'{payload.calidad_aceite_transmision_bayoneta_observacion}' , \
        @calidad_aceite_transmision_bayoneta_foto = N'{payload.calidad_aceite_transmision_bayoneta_foto}' , \
        @liquido_bateria_condiciones = {payload.liquido_bateria_condiciones}, \
        @liquido_bateria_condiciones_observacion = N'{payload.liquido_bateria_condiciones_observacion}' , \
        @liquido_bateria_condiciones_foto = N'{payload.liquido_bateria_condiciones_foto}' ,    \
        @bandas_poly_v = {payload.bandas_poly_v}, \
        @bandas_poly_v_observacion = N'{payload.bandas_poly_v_observacion}' , \
        @bandas_poly_v_foto = N'{payload.bandas_poly_v_foto}' , \
        @poleas_banda = {payload.poleas_banda}, \
        @poleas_banda_observacion = N'{payload.poleas_banda_observacion}' , \
        @poleas_banda_foto = N'{payload.poleas_banda_foto}' , \
        @banda_tiempo = {payload.banda_tiempo}, \
        @banda_tiempo_observacion = N'{payload.banda_tiempo_observacion}' , \
        @banda_tiempo_foto = N'{payload.banda_tiempo_foto}' , \
        @otros_habitaculo_motor = {payload.otros_habitaculo_motor}, \
        @otros_habitaculo_motor_observacion = N'{payload.otros_habitaculo_motor_observacion}' , \
        @otros_habitaculo_motor_foto = N'{payload.otros_habitaculo_motor_foto}' , \
        @reset_intervalo_servicio = {payload.reset_intervalo_servicio}, \
        @reset_intervalo_servicio_observacion = N'{payload.reset_intervalo_servicio_observacion}' , \
        @reset_intervalo_servicio_foto = N'{payload.reset_intervalo_servicio_foto}' , \
        @ajuste_tornillos_neumaticos_torquimetro = {payload.ajuste_tornillos_neumaticos_torquimetro}, \
        @ajuste_tornillos_neumaticos_torquimetro_observacion = N'{payload.ajuste_tornillos_neumaticos_torquimetro_observacion}' , \
        @ajuste_tornillos_neumaticos_torquimetro_foto = N'{payload.ajuste_tornillos_neumaticos_torquimetro_foto}' , \
        @limpiar_libricar_puertas_cerraduras = {payload.limpiar_libricar_puertas_cerraduras}, \
        @limpiar_libricar_puertas_cerraduras_observacion = N'{payload.limpiar_libricar_puertas_cerraduras_observacion}' , \
        @limpiar_libricar_puertas_cerraduras_foto = N'{payload.limpiar_libricar_puertas_cerraduras_foto}' , \
        @completar_plan_mantenimiento = {payload.completar_plan_mantenimiento}, \
        @completar_plan_mantenimiento_observacion = N'{payload.completar_plan_mantenimiento_observacion}' , \
        @completar_plan_mantenimiento_foto = N'{payload.completar_plan_mantenimiento_foto}' , \
        @fecha = N'{payload.Fecha}' , \
        @IdEmpleado = {payload.IdEmpleado} , \
        @IdVehiculo = {payload.IdVehiculo} , \
        @Id_ordendeservicio = {payload.Id_ordendeservicio} """
#modificaciones adjuntas pruebas
    try:
        with engine.begin() as conn:
            # Ejecutar la consulta SQL
            roles_df = pd.read_sql(query, conn)

        # Crear el objeto de respuesta
        response = ResponseModel(id_resultado=1, respuesta="El checklist se actualizó correctamente.")
        return JSONResponse(status_code=200, content=response.dict())

    except Exception as e:
        # Manejo de errores
        raise HTTPException(status_code=500, detail=str(e))

@app.post(
        path="/api/servicio",
        name='Insertar servicio',
        tags=['Servicio'],
        description='Método para insertar el servicio',
        response_model=Checklist
)
def saveservicio(payload: Checklist):
    query = f"""
    EXEC InsertServicio
        @lectura_codigos = {payload.lectura_codigos}, \
        @lectura_codigos_observacion = N'{payload.lectura_codigos_observacion}' , \
        @lectura_codigos_foto = N'{payload.lectura_codigos_foto}' , \
        @servofreno = {payload.servofreno}, \
        @servofreno_observacion = N'{payload.servofreno_observacion}' , \
        @servofreno_foto = N'{payload.servofreno_foto}' , \
        @pedal_freno = {payload.pedal_freno}, \
        @pedal_freno_observacion = N'{payload.pedal_freno_observacion}' , \
        @pedal_freno_foto = N'{payload.pedal_freno_foto}' , \
        @pedal_estacionamiento = {payload.pedal_estacionamiento}, \
        @pedal_estacionamiento_observacion = N'{payload.pedal_estacionamiento_observacion}' , \
        @pedal_estacionamiento_foto = N'{payload.pedal_estacionamiento_foto}' , \
        @cinturon_seguridad = {payload.cinturon_seguridad}, \
        @cinturon_seguridad_observacion = N'{payload.cinturon_seguridad_observacion}' , \
        @cinturon_seguridad_foto = N'{payload.cinturon_seguridad_foto}' , \
        @cuadro_instrumentos = {payload.cuadro_instrumentos}, \
        @cuadro_instrumentos_observacion = N'{payload.cuadro_instrumentos_observacion}' ,\
        @cuadro_instrumentos_foto = N'{payload.cuadro_instrumentos_foto}' ,\
        @aire_acondicionado = {payload.aire_acondicionado},\
        @aire_acondicionado_observacion = N'{payload.aire_acondicionado_observacion}' ,\
        @aire_acondicionado_foto = N'{payload.aire_acondicionado_foto}' ,\
        @bocina_claxon = {payload.bocina_claxon},\
        @bocina_claxon_observacion = N'{payload.bocina_claxon_observacion}' ,\
        @bocina_claxon_foto = N'{payload.bocina_claxon_foto}' ,\
        @iluminacion_interior = {payload.iluminacion_interior},\
        @iluminacion_interior_observacion = N'{payload.iluminacion_interior_observacion}' ,\
        @iluminacion_interior_foto = N'{payload.iluminacion_interior_foto}' ,\
        @iluminacion_externa = {payload.iluminacion_externa},\
        @iluminacion_externa_observacion = N'{payload.iluminacion_externa_observacion}' ,\
        @iluminacion_externa_foto = N'{payload.iluminacion_externa_foto}' ,\
        @limpiaparabrisas = {payload.limpiaparabrisas}, \
        @limpiaparabrisas_observacion = N'{payload.limpiaparabrisas_observacion}' ,\
        @limpiaparabrisas_foto = N'{payload.limpiaparabrisas_foto}' , \
        @limpia_medallon = {payload.limpia_medallon}, \
        @limpia_medallon_observacion = N'{payload.limpia_medallon_observacion}' , \
        @limpia_medallon_foto = N'{payload.limpia_medallon_foto}' , \
        @neumaticos_friccion = {payload.neumaticos_friccion}, \
        @neumaticos_friccion_observacion = N'{payload.neumaticos_friccion_observacion}' ,  \
        @neumaticos_friccion_foto = N'{payload.neumaticos_friccion_foto}' , \
        @otros_vehiculo_en_piso = {payload.otros_vehiculo_en_piso}, \
        @otros_vehiculo_en_piso_observacion = N'{payload.otros_vehiculo_en_piso_observacion}' , \
        @otros_vehiculo_en_piso_foto = N'{payload.otros_vehiculo_en_piso_foto}' , \
        @estado_fugas_aceite = {payload.estado_fugas_aceite}, \
        @estado_fugas_aceite_observacion = N'{payload.estado_fugas_aceite_observacion}' , \
        @estado_fugas_aceite_foto = N'{payload.estado_fugas_aceite_foto}' , \
        @estado_nivel_calidad_lubricante_transmision = {payload.estado_nivel_calidad_lubricante_transmision}, \
        @estado_nivel_calidad_lubricante_transmision_observacion = N'{payload.estado_nivel_calidad_lubricante_transmision_observacion}' , \
        @estado_nivel_calidad_lubricante_transmision_foto = N'{payload.estado_nivel_calidad_lubricante_transmision_foto}' , \
        @estado_nivel_calidad_lubricante_diferencial = {payload.estado_nivel_calidad_lubricante_diferencial}, \
        @estado_nivel_calidad_lubricante_diferencial_observacion = N'{payload.estado_nivel_calidad_lubricante_diferencial_observacion}' , \
        @estado_nivel_calidad_lubricante_diferencial_foto = N'{payload.estado_nivel_calidad_lubricante_diferencial_foto}' , \
        @cubrepolvos_flechas = {payload.cubrepolvos_flechas}, \
        @cubrepolvos_flechas_observacion = N'{payload.cubrepolvos_flechas_observacion}' , \
        @cubrepolvos_flechas_foto = N'{payload.cubrepolvos_flechas_foto}' , \
        @componentes_direccion = {payload.componentes_direccion}, \
        @componentes_direccion_observacion = N'{payload.componentes_direccion_observacion}' , \
        @componentes_direccion_foto = N'{payload.componentes_direccion_foto}' , \
        @componentes_suspesion = {payload.componentes_suspesion}, \
        @componentes_suspesion_observacion = N'{payload.componentes_suspesion_observacion}' , \
        @componentes_suspesion_foto = N'{payload.componentes_suspesion_foto}' , \
        @sistema_escape_completo = {payload.sistema_escape_completo}, \
        @sistema_escape_completo_observacion = N'{payload.sistema_escape_completo_observacion}' , \
        @sistema_escape_completo_foto = N'{payload.sistema_escape_completo_foto}' , \
        @sistema_alimentacion_combustible = {payload.sistema_alimentacion_combustible}, \
        @sistema_alimentacion_combustible_observacion = N'{payload.sistema_alimentacion_combustible_observacion}' , \
        @sistema_alimentacion_combustible_foto = N'{payload.sistema_alimentacion_combustible_foto}' , \
        @filtro_combustible = {payload.filtro_combustible}, \
        @filtro_combustible_observacion = N'{payload.filtro_combustible_observacion}' , \
        @filtro_combustible_foto = N'{payload.filtro_combustible_foto}' , \
        @control_fugas_direccion_hidraulica = {payload.control_fugas_direccion_hidraulica}, \
        @control_fugas_direccion_hidraulica_observacion = N'{payload.control_fugas_direccion_hidraulica_observacion}' , \
        @control_fugas_direccion_hidraulica_foto = N'{payload.control_fugas_direccion_hidraulica_foto}' , \
        @otros_altura_total = {payload.otros_altura_total}, \
        @otros_altura_total_observacion = N'{payload.otros_altura_total_observacion}' , \
        @otros_altura_total_foto = N'{payload.otros_altura_total_foto}' , \
        @rodamiento_mazas_rueda = {payload.rodamiento_mazas_rueda}, \
        @rodamiento_mazas_rueda_observacion = N'{payload.rodamiento_mazas_rueda_observacion}' , \
        @rodamiento_mazas_rueda_foto = N'{payload.rodamiento_mazas_rueda_foto}' , \
        @holgura_partes_suspension_rueda = {payload.holgura_partes_suspension_rueda}, \
        @holgura_partes_suspension_rueda_observacion = N'{payload.holgura_partes_suspension_rueda_observacion}' , \
        @holgura_partes_suspension_rueda_foto = N'{payload.holgura_partes_suspension_rueda_foto}' , \
        @control_neumaticos_desgaste_presion = {payload.control_neumaticos_desgaste_presion}, \
        @control_neumaticos_desgaste_presion_observacion = N'{payload.control_neumaticos_desgaste_presion_observacion}' , \
        @control_neumaticos_desgaste_presion_foto = N'{payload.control_neumaticos_desgaste_presion_foto}' , \
        @profundidad = {payload.profundidad}, \
        @profundidad_observacion = N'{payload.profundidad_observacion}' , \
        @profundidad_foto = N'{payload.profundidad_foto}' , \
        @presion = {payload.presion}, \
        @presion_observacion = N'{payload.presion_observacion}' , \
        @presion_foto = N'{payload.presion_foto}' , \
        @otros_altura_media = {payload.otros_altura_media}, \
        @otros_altura_media_observacion = N'{payload.otros_altura_media_observacion}' , \
        @otros_altura_media_foto = N'{payload.otros_altura_media_foto}' , \
        @nivel_calidad_aceite_motor = {payload.nivel_calidad_aceite_motor}, \
        @nivel_calidad_aceite_motor_observacion = N'{payload.nivel_calidad_aceite_motor_observacion}' , \
        @nivel_calidad_aceite_motor_foto = N'{payload.nivel_calidad_aceite_motor_foto}' , \
        @filtro_aire = {payload.filtro_aire}, \
        @filtro_aire_observacion = N'{payload.filtro_aire_observacion}' , \
        @filtro_aire_foto = N'{payload.filtro_aire_foto}' , \
        @filtro_polen = {payload.filtro_polen}, \
        @filtro_polen_observacion = N'{payload.filtro_polen_observacion}' , \
        @filtro_polen_foto = N'{payload.filtro_polen_foto}' , \
        @filtro_pcv = {payload.filtro_pcv}, \
        @filtro_pcv_observacion = N'{payload.filtro_pcv_observacion}' , \
        @filtro_pcv_foto = N'{payload.filtro_pcv_foto}' , \
        @valvula_pcv = {payload.valvula_pcv}, \
        @valvula_pcv_observacion = N'{payload.valvula_pcv_observacion}' , \
        @valvula_pcv_foto = N'{payload.valvula_pcv_foto}' , \
        @bujias_encendido = {payload.bujias_encendido}, \
        @bujias_encendido_observacion = N'{payload.bujias_encendido_observacion}' , \
        @bujias_encendido_foto = N'{payload.bujias_encendido_foto}' , \
        @cables_bujias_bobinas_ignicion = {payload.cables_bujias_bobinas_ignicion}, \
        @cables_bujias_bobinas_ignicion_observacion = N'{payload.cables_bujias_bobinas_ignicion_observacion}' , \
        @cables_bujias_bobinas_ignicion_foto = N'{payload.cables_bujias_bobinas_ignicion_foto}' , \
        @nivel_anticongenlante = {payload.nivel_anticongenlante}, \
        @nivel_anticongenlante_observacion = N'{payload.nivel_anticongenlante_observacion}' , \
        @nivel_anticongenlante_foto = N'{payload.nivel_anticongenlante_foto}' , \
        @tapon_radiador = {payload.tapon_radiador}, \
        @tapon_radiador_observacion = N'{payload.tapon_radiador_observacion}' , \
        @tapon_radiador_foto = N'{payload.tapon_radiador_foto}' , \
        @mangueras_sistema = {payload.mangueras_sistema}, \
        @mangueras_sistema_observacion = N'{payload.mangueras_sistema_observacion}' , \
        @mangueras_sistema_foto = N'{payload.mangueras_sistema_foto}' , \
        @desempeño_ventilador = {payload.desempeño_ventilador}, \
        @desempeño_ventilador_observacion = N'{payload.desempeño_ventilador_observacion}' , \
        @desempeño_ventilador_foto = N'{payload.desempeño_ventilador_foto}' , \
        @calidad_liquido_limpiaparabrisas = {payload.calidad_liquido_limpiaparabrisas}, \
        @calidad_liquido_limpiaparabrisas_observacion = N'{payload.calidad_liquido_limpiaparabrisas_observacion}' , \
        @calidad_liquido_limpiaparabrisas_foto = N'{payload.calidad_liquido_limpiaparabrisas_foto}' , \
        @calidad_aceite_direccion_hidraulica = {payload.calidad_aceite_direccion_hidraulica}, \
        @calidad_aceite_direccion_hidraulica_observacion = N'{payload.calidad_aceite_direccion_hidraulica_observacion}' , \
        @calidad_aceite_direccion_hidraulica_foto = N'{payload.calidad_aceite_direccion_hidraulica_foto}' , \
        @calidad_aceite_transmision_bayoneta = {payload.calidad_aceite_transmision_bayoneta}, \
        @calidad_aceite_transmision_bayoneta_observacion = N'{payload.calidad_aceite_transmision_bayoneta_observacion}' , \
        @calidad_aceite_transmision_bayoneta_foto = N'{payload.calidad_aceite_transmision_bayoneta_foto}' , \
        @liquido_bateria_condiciones = {payload.liquido_bateria_condiciones}, \
        @liquido_bateria_condiciones_observacion = N'{payload.liquido_bateria_condiciones_observacion}' , \
        @liquido_bateria_condiciones_foto = N'{payload.liquido_bateria_condiciones_foto}' ,    \
        @bandas_poly_v = {payload.bandas_poly_v}, \
        @bandas_poly_v_observacion = N'{payload.bandas_poly_v_observacion}' , \
        @bandas_poly_v_foto = N'{payload.bandas_poly_v_foto}' , \
        @poleas_banda = {payload.poleas_banda}, \
        @poleas_banda_observacion = N'{payload.poleas_banda_observacion}' , \
        @poleas_banda_foto = N'{payload.poleas_banda_foto}' , \
        @banda_tiempo = {payload.banda_tiempo}, \
        @banda_tiempo_observacion = N'{payload.banda_tiempo_observacion}' , \
        @banda_tiempo_foto = N'{payload.banda_tiempo_foto}' , \
        @otros_habitaculo_motor = {payload.otros_habitaculo_motor}, \
        @otros_habitaculo_motor_observacion = N'{payload.otros_habitaculo_motor_observacion}' , \
        @otros_habitaculo_motor_foto = N'{payload.otros_habitaculo_motor_foto}' , \
        @reset_intervalo_servicio = {payload.reset_intervalo_servicio}, \
        @reset_intervalo_servicio_observacion = N'{payload.reset_intervalo_servicio_observacion}' , \
        @reset_intervalo_servicio_foto = N'{payload.reset_intervalo_servicio_foto}' , \
        @ajuste_tornillos_neumaticos_torquimetro = {payload.ajuste_tornillos_neumaticos_torquimetro}, \
        @ajuste_tornillos_neumaticos_torquimetro_observacion = N'{payload.ajuste_tornillos_neumaticos_torquimetro_observacion}' , \
        @ajuste_tornillos_neumaticos_torquimetro_foto = N'{payload.ajuste_tornillos_neumaticos_torquimetro_foto}' , \
        @limpiar_libricar_puertas_cerraduras = {payload.limpiar_libricar_puertas_cerraduras}, \
        @limpiar_libricar_puertas_cerraduras_observacion = N'{payload.limpiar_libricar_puertas_cerraduras_observacion}' , \
        @limpiar_libricar_puertas_cerraduras_foto = N'{payload.limpiar_libricar_puertas_cerraduras_foto}' , \
        @completar_plan_mantenimiento = {payload.completar_plan_mantenimiento}, \
        @completar_plan_mantenimiento_observacion = N'{payload.completar_plan_mantenimiento_observacion}' , \
        @completar_plan_mantenimiento_foto = N'{payload.completar_plan_mantenimiento_foto}' , \
        @fecha = N'{payload.Fecha}' , \
        @Id_empleado = {payload.IdEmpleado} , \
        @Id_vehiculo = {payload.IdVehiculo} , \
        @id_ordendeservicio = {payload.Id_ordendeservicio} , \
        @id_checklist = {payload.id_checklist} , \
        @NumeroSerie = '{payload.NumeroSerie}', \
        @Activo = {payload.Activo}"""
    print (query)
    with engine.begin() as conn:
          conn.execution_options(autocommit = True)
          roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="El servicio se guardó de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)


@app.put(
    path="/api/servicio",
    name="Actualizar Servicio",
    tags=["Servicio"],
    description="Método para actualizar la información del servicio",
    response_model=Checklist
)
def put_servicio(payload: Checklist):
    # Construcción de la consulta SQL para ejecutar el procedimiento almacenado
    query = f"""
    exec [dbo].[UpdateServicio]
        @lectura_codigos = {payload.lectura_codigos}, \
        @lectura_codigos_observacion = N'{payload.lectura_codigos_observacion}' , \
        @lectura_codigos_foto = N'{payload.lectura_codigos_foto}' , \
        @servofreno = {payload.servofreno}, \
        @servofreno_observacion = N'{payload.servofreno_observacion}' , \
        @servofreno_foto = N'{payload.servofreno_foto}' , \
        @pedal_freno = {payload.pedal_freno}, \
        @pedal_freno_observacion = N'{payload.pedal_freno_observacion}' , \
        @pedal_freno_foto = N'{payload.pedal_freno_foto}' , \
        @pedal_estacionamiento = {payload.pedal_estacionamiento}, \
        @pedal_estacionamiento_observacion = N'{payload.pedal_estacionamiento_observacion}' , \
        @pedal_estacionamiento_foto = N'{payload.pedal_estacionamiento_foto}' , \
        @cinturon_seguridad = {payload.cinturon_seguridad}, \
        @cinturon_seguridad_observacion = N'{payload.cinturon_seguridad_observacion}' , \
        @cinturon_seguridad_foto = N'{payload.cinturon_seguridad_foto}' , \
        @cuadro_instrumentos = {payload.cuadro_instrumentos}, \
        @cuadro_instrumentos_observacion = N'{payload.cuadro_instrumentos_observacion}' ,\
        @cuadro_instrumentos_foto = N'{payload.cuadro_instrumentos_foto}' ,\
        @aire_acondicionado = {payload.aire_acondicionado},\
        @aire_acondicionado_observacion = N'{payload.aire_acondicionado_observacion}' ,\
        @aire_acondicionado_foto = N'{payload.aire_acondicionado_foto}' ,\
        @bocina_claxon = {payload.bocina_claxon},\
        @bocina_claxon_observacion = N'{payload.bocina_claxon_observacion}' ,\
        @bocina_claxon_foto = N'{payload.bocina_claxon_foto}' ,\
        @iluminacion_interior = {payload.iluminacion_interior},\
        @iluminacion_interior_observacion = N'{payload.iluminacion_interior_observacion}' ,\
        @iluminacion_interior_foto = N'{payload.iluminacion_interior_foto}' ,\
        @iluminacion_externa = {payload.iluminacion_externa},\
        @iluminacion_externa_observacion = N'{payload.iluminacion_externa_observacion}' ,\
        @iluminacion_externa_foto = N'{payload.iluminacion_externa_foto}' ,\
        @limpiaparabrisas = {payload.limpiaparabrisas}, \
        @limpiaparabrisas_observacion = N'{payload.limpiaparabrisas_observacion}' ,\
        @limpiaparabrisas_foto = N'{payload.limpiaparabrisas_foto}' , \
        @limpia_medallon = {payload.limpia_medallon}, \
        @limpia_medallon_observacion = N'{payload.limpia_medallon_observacion}' , \
        @limpia_medallon_foto = N'{payload.limpia_medallon_foto}' , \
        @neumaticos_friccion = {payload.neumaticos_friccion}, \
        @neumaticos_friccion_observacion = N'{payload.neumaticos_friccion_observacion}' ,  \
        @neumaticos_friccion_foto = N'{payload.neumaticos_friccion_foto}' , \
        @otros_vehiculo_en_piso = {payload.otros_vehiculo_en_piso}, \
        @otros_vehiculo_en_piso_observacion = N'{payload.otros_vehiculo_en_piso_observacion}' , \
        @otros_vehiculo_en_piso_foto = N'{payload.otros_vehiculo_en_piso_foto}' , \
        @estado_fugas_aceite = {payload.estado_fugas_aceite}, \
        @estado_fugas_aceite_observacion = N'{payload.estado_fugas_aceite_observacion}' , \
        @estado_fugas_aceite_foto = N'{payload.estado_fugas_aceite_foto}' , \
        @estado_nivel_calidad_lubricante_transmision = {payload.estado_nivel_calidad_lubricante_transmision}, \
        @estado_nivel_calidad_lubricante_transmision_observacion = N'{payload.estado_nivel_calidad_lubricante_transmision_observacion}' , \
        @estado_nivel_calidad_lubricante_transmision_foto = N'{payload.estado_nivel_calidad_lubricante_transmision_foto}' , \
        @estado_nivel_calidad_lubricante_diferencial = {payload.estado_nivel_calidad_lubricante_diferencial}, \
        @estado_nivel_calidad_lubricante_diferencial_observacion = N'{payload.estado_nivel_calidad_lubricante_diferencial_observacion}' , \
        @estado_nivel_calidad_lubricante_diferencial_foto = N'{payload.estado_nivel_calidad_lubricante_diferencial_foto}' , \
        @cubrepolvos_flechas = {payload.cubrepolvos_flechas}, \
        @cubrepolvos_flechas_observacion = N'{payload.cubrepolvos_flechas_observacion}' , \
        @cubrepolvos_flechas_foto = N'{payload.cubrepolvos_flechas_foto}' , \
        @componentes_direccion = {payload.componentes_direccion}, \
        @componentes_direccion_observacion = N'{payload.componentes_direccion_observacion}' , \
        @componentes_direccion_foto = N'{payload.componentes_direccion_foto}' , \
        @componentes_suspesion = {payload.componentes_suspesion}, \
        @componentes_suspesion_observacion = N'{payload.componentes_suspesion_observacion}' , \
        @componentes_suspesion_foto = N'{payload.componentes_suspesion_foto}' , \
        @sistema_escape_completo = {payload.sistema_escape_completo}, \
        @sistema_escape_completo_observacion = N'{payload.sistema_escape_completo_observacion}' , \
        @sistema_escape_completo_foto = N'{payload.sistema_escape_completo_foto}' , \
        @sistema_alimentacion_combustible = {payload.sistema_alimentacion_combustible}, \
        @sistema_alimentacion_combustible_observacion = N'{payload.sistema_alimentacion_combustible_observacion}' , \
        @sistema_alimentacion_combustible_foto = N'{payload.sistema_alimentacion_combustible_foto}' , \
        @filtro_combustible = {payload.filtro_combustible}, \
        @filtro_combustible_observacion = N'{payload.filtro_combustible_observacion}' , \
        @filtro_combustible_foto = N'{payload.filtro_combustible_foto}' , \
        @control_fugas_direccion_hidraulica = {payload.control_fugas_direccion_hidraulica}, \
        @control_fugas_direccion_hidraulica_observacion = N'{payload.control_fugas_direccion_hidraulica_observacion}' , \
        @control_fugas_direccion_hidraulica_foto = N'{payload.control_fugas_direccion_hidraulica_foto}' , \
        @otros_altura_total = {payload.otros_altura_total}, \
        @otros_altura_total_observacion = N'{payload.otros_altura_total_observacion}' , \
        @otros_altura_total_foto = N'{payload.otros_altura_total_foto}' , \
        @rodamiento_mazas_rueda = {payload.rodamiento_mazas_rueda}, \
        @rodamiento_mazas_rueda_observacion = N'{payload.rodamiento_mazas_rueda_observacion}' , \
        @rodamiento_mazas_rueda_foto = N'{payload.rodamiento_mazas_rueda_foto}' , \
        @holgura_partes_suspension_rueda = {payload.holgura_partes_suspension_rueda}, \
        @holgura_partes_suspension_rueda_observacion = N'{payload.holgura_partes_suspension_rueda_observacion}' , \
        @holgura_partes_suspension_rueda_foto = N'{payload.holgura_partes_suspension_rueda_foto}' , \
        @control_neumaticos_desgaste_presion = {payload.control_neumaticos_desgaste_presion}, \
        @control_neumaticos_desgaste_presion_observacion = N'{payload.control_neumaticos_desgaste_presion_observacion}' , \
        @control_neumaticos_desgaste_presion_foto = N'{payload.control_neumaticos_desgaste_presion_foto}' , \
        @profundidad = {payload.profundidad}, \
        @profundidad_observacion = N'{payload.profundidad_observacion}' , \
        @profundidad_foto = N'{payload.profundidad_foto}' , \
        @presion = {payload.presion}, \
        @presion_observacion = N'{payload.presion_observacion}' , \
        @presion_foto = N'{payload.presion_foto}' , \
        @otros_altura_media = {payload.otros_altura_media}, \
        @otros_altura_media_observacion = N'{payload.otros_altura_media_observacion}' , \
        @otros_altura_media_foto = N'{payload.otros_altura_media_foto}' , \
        @nivel_calidad_aceite_motor = {payload.nivel_calidad_aceite_motor}, \
        @nivel_calidad_aceite_motor_observacion = N'{payload.nivel_calidad_aceite_motor_observacion}' , \
        @nivel_calidad_aceite_motor_foto = N'{payload.nivel_calidad_aceite_motor_foto}' , \
        @filtro_aire = {payload.filtro_aire}, \
        @filtro_aire_observacion = N'{payload.filtro_aire_observacion}' , \
        @filtro_aire_foto = N'{payload.filtro_aire_foto}' , \
        @filtro_polen = {payload.filtro_polen}, \
        @filtro_polen_observacion = N'{payload.filtro_polen_observacion}' , \
        @filtro_polen_foto = N'{payload.filtro_polen_foto}' , \
        @filtro_pcv = {payload.filtro_pcv}, \
        @filtro_pcv_observacion = N'{payload.filtro_pcv_observacion}' , \
        @filtro_pcv_foto = N'{payload.filtro_pcv_foto}' , \
        @valvula_pcv = {payload.valvula_pcv}, \
        @valvula_pcv_observacion = N'{payload.valvula_pcv_observacion}' , \
        @valvula_pcv_foto = N'{payload.valvula_pcv_foto}' , \
        @bujias_encendido = {payload.bujias_encendido}, \
        @bujias_encendido_observacion = N'{payload.bujias_encendido_observacion}' , \
        @bujias_encendido_foto = N'{payload.bujias_encendido_foto}' , \
        @cables_bujias_bobinas_ignicion = {payload.cables_bujias_bobinas_ignicion}, \
        @cables_bujias_bobinas_ignicion_observacion = N'{payload.cables_bujias_bobinas_ignicion_observacion}' , \
        @cables_bujias_bobinas_ignicion_foto = N'{payload.cables_bujias_bobinas_ignicion_foto}' , \
        @nivel_anticongenlante = {payload.nivel_anticongenlante}, \
        @nivel_anticongenlante_observacion = N'{payload.nivel_anticongenlante_observacion}' , \
        @nivel_anticongenlante_foto = N'{payload.nivel_anticongenlante_foto}' , \
        @tapon_radiador = {payload.tapon_radiador}, \
        @tapon_radiador_observacion = N'{payload.tapon_radiador_observacion}' , \
        @tapon_radiador_foto = N'{payload.tapon_radiador_foto}' , \
        @mangueras_sistema = {payload.mangueras_sistema}, \
        @mangueras_sistema_observacion = N'{payload.mangueras_sistema_observacion}' , \
        @mangueras_sistema_foto = N'{payload.mangueras_sistema_foto}' , \
        @desempeño_ventilador = {payload.desempeño_ventilador}, \
        @desempeño_ventilador_observacion = N'{payload.desempeño_ventilador_observacion}' , \
        @desempeño_ventilador_foto = N'{payload.desempeño_ventilador_foto}' , \
        @calidad_liquido_limpiaparabrisas = {payload.calidad_liquido_limpiaparabrisas}, \
        @calidad_liquido_limpiaparabrisas_observacion = N'{payload.calidad_liquido_limpiaparabrisas_observacion}' , \
        @calidad_liquido_limpiaparabrisas_foto = N'{payload.calidad_liquido_limpiaparabrisas_foto}' , \
        @calidad_aceite_direccion_hidraulica = {payload.calidad_aceite_direccion_hidraulica}, \
        @calidad_aceite_direccion_hidraulica_observacion = N'{payload.calidad_aceite_direccion_hidraulica_observacion}' , \
        @calidad_aceite_direccion_hidraulica_foto = N'{payload.calidad_aceite_direccion_hidraulica_foto}' , \
        @calidad_aceite_transmision_bayoneta = {payload.calidad_aceite_transmision_bayoneta}, \
        @calidad_aceite_transmision_bayoneta_observacion = N'{payload.calidad_aceite_transmision_bayoneta_observacion}' , \
        @calidad_aceite_transmision_bayoneta_foto = N'{payload.calidad_aceite_transmision_bayoneta_foto}' , \
        @liquido_bateria_condiciones = {payload.liquido_bateria_condiciones}, \
        @liquido_bateria_condiciones_observacion = N'{payload.liquido_bateria_condiciones_observacion}' , \
        @liquido_bateria_condiciones_foto = N'{payload.liquido_bateria_condiciones_foto}' ,    \
        @bandas_poly_v = {payload.bandas_poly_v}, \
        @bandas_poly_v_observacion = N'{payload.bandas_poly_v_observacion}' , \
        @bandas_poly_v_foto = N'{payload.bandas_poly_v_foto}' , \
        @poleas_banda = {payload.poleas_banda}, \
        @poleas_banda_observacion = N'{payload.poleas_banda_observacion}' , \
        @poleas_banda_foto = N'{payload.poleas_banda_foto}' , \
        @banda_tiempo = {payload.banda_tiempo}, \
        @banda_tiempo_observacion = N'{payload.banda_tiempo_observacion}' , \
        @banda_tiempo_foto = N'{payload.banda_tiempo_foto}' , \
        @otros_habitaculo_motor = {payload.otros_habitaculo_motor}, \
        @otros_habitaculo_motor_observacion = N'{payload.otros_habitaculo_motor_observacion}' , \
        @otros_habitaculo_motor_foto = N'{payload.otros_habitaculo_motor_foto}' , \
        @reset_intervalo_servicio = {payload.reset_intervalo_servicio}, \
        @reset_intervalo_servicio_observacion = N'{payload.reset_intervalo_servicio_observacion}' , \
        @reset_intervalo_servicio_foto = N'{payload.reset_intervalo_servicio_foto}' , \
        @ajuste_tornillos_neumaticos_torquimetro = {payload.ajuste_tornillos_neumaticos_torquimetro}, \
        @ajuste_tornillos_neumaticos_torquimetro_observacion = N'{payload.ajuste_tornillos_neumaticos_torquimetro_observacion}' , \
        @ajuste_tornillos_neumaticos_torquimetro_foto = N'{payload.ajuste_tornillos_neumaticos_torquimetro_foto}' , \
        @limpiar_libricar_puertas_cerraduras = {payload.limpiar_libricar_puertas_cerraduras}, \
        @limpiar_libricar_puertas_cerraduras_observacion = N'{payload.limpiar_libricar_puertas_cerraduras_observacion}' , \
        @limpiar_libricar_puertas_cerraduras_foto = N'{payload.limpiar_libricar_puertas_cerraduras_foto}' , \
        @completar_plan_mantenimiento = {payload.completar_plan_mantenimiento}, \
        @completar_plan_mantenimiento_observacion = N'{payload.completar_plan_mantenimiento_observacion}' , \
        @completar_plan_mantenimiento_foto = N'{payload.completar_plan_mantenimiento_foto}' , \
        @fecha = N'{payload.Fecha}' , \
        @IdEmpleado = {payload.IdEmpleado} , \
        @IdVehiculo = {payload.IdVehiculo} , \
        @Id_ordendeservicio = {payload.Id_ordendeservicio} """
#modificaciones adjuntas pruebas
    try:
        with engine.begin() as conn:
            # Ejecutar la consulta SQL
            roles_df = pd.read_sql(query, conn)

        # Crear el objeto de respuesta
        response = ResponseModel(id_resultado=1, respuesta="El servicio se actualizó correctamente.")
        return JSONResponse(status_code=200, content=response.dict())

    except Exception as e:
        # Manejo de errores
        raise HTTPException(status_code=500, detail=str(e))




@app.get(
        path="/api/obtenerservicio",
        name='Obtener servicio',
        tags=['Servicio'],
        description='Método para obtener la informacion de 1 checklist',
        response_model=Checklist
)
def getservicio(Idchecklist: int):
    query = f"exec [dbo].[sp_get_all_servicio] @IdCheckList = {Idchecklist}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])


@app.get(
        path="/api/obtenerflotillaporid",
        name='Obtener flotilla por id',
        tags=['Flotillas'],
        description='Método para obtener la informacion de 1 flotilla',
        response_model=Flotillas
)
def getflotillaevidencia(IdFlotilla: int):
    query = f"exec [dbo].[ObtenerAllFlotillasPorID] @IdFlotilla = {IdFlotilla}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])

@app.get(
        path="/api/obtenerallflotillas",
        name='Obtener todas las flotillas',
        tags=['Flotillas'],
        description='Método para obtener la informacion de todos las flotillas',
        response_model=Flotillas
)
def getsallflotillas():
    query = f"exec [dbo].[ObtenerAllFlotillas]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)


@app.get(
        path="/api/obtenerhistoricos",
        name='Obtener históricos',
        tags=['Historico'],
        description='Método para obtener la informacion de todos los históricos',
        response_model=CheckListHistorico
)
def getshistoricocheck():
    query = f"exec [dbo].[ObtenerHistoricoCheck]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/obtenerhistoricosservicios",
        name='Obtener históricos servicios',
        tags=['Historico'],
        description='Método para obtener la informacion de todos los históricos de servicios',
        response_model=CheckListHistorico
)
def getshistoricoservicio():
    query = f"exec [dbo].[ObtenerHistoricoServicio]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/obtenerflotillas",
        name='Obtener flotillas',
        tags=['Flotillas'],
        description='Método para obtener la informacion de todas las flotillas',
        response_model=Flotillas
)
def getsflotillas():
    query = f"exec [dbo].[ObtenerFlotillas]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/obtenerflotilla",
        name='Obtener flotilla',
        tags=['Flotillas'],
        description='Método para obtener la informacion de 1 flotilla',
        response_model=Flotillas
)
def getflotilla(IdFlotilla: int):
    query = f"exec [dbo].[sp_get_all_flotillas] @IdFlotilla = {IdFlotilla}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])

@app.post(
    path="/api/flotilla",
    name="Guarda Flotilla",
    tags=["Flotillas"],
    description="Método para guardar la información de las flotillas",
    response_model=Flotillas,
)
def guardarFlotilla(payload: Flotillas):
    try:
        # Crear el diccionario de parámetros desde el payload
        parametros = payload.dict()

        # Definir la consulta SQL para ejecutar el procedimiento almacenado
        query = text("""
            EXEC dbo.Insertflotillas
                @NamesFlotillas = :NamesFlotillas,
                @Encargado = :Encargado
        """)

        # Ejecutar la consulta SQL
        with engine.begin() as conn:
            conn.execute(query, parametros)

        # Respuesta de éxito
        return JSONResponse(
            status_code=200,
            content={
                "id_resultado": 1,
                "respuesta": "La flotilla se guardó de manera correcta",
                "detalles": parametros,
            },
        )

    except Exception as e:
        # Respuesta de error
        raise HTTPException(
            status_code=500,
            detail=f"Error al guardar la flotilla: {str(e)}",
        )
@app.get(
        path="/api/obtenerservicios",
        name='Obtener servicios',
        tags=['Servicio'],
        description='Método para obtener la informacion todos los servicios',
        response_model=Checklist
)
def getservicios(IdEmpleado:int):
    query = f"exec [dbo].[ObtenerServicios] @IdEmpleado = {IdEmpleado}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.put(
    path="/api/checklist",
    name="Actualizar checklist",
    tags=["Checklist"],
    description="Método para actualizar la información del checklist",
    response_model=Checklist
)
def put_servicios(payload: Checklist):
    # Construcción de la consulta SQL para ejecutar el procedimiento almacenado
    query = f"""
    exec [dbo].[UpdateServicio]
                @lectura_codigos = {payload.lectura_codigos}, \
        @lectura_codigos_observacion = N'{payload.lectura_codigos_observacion}' , \
        @lectura_codigos_foto = N'{payload.lectura_codigos_foto}' , \
        @servofreno = {payload.servofreno}, \
        @servofreno_observacion = N'{payload.servofreno_observacion}' , \
        @servofreno_foto = N'{payload.servofreno_foto}' , \
        @pedal_freno = {payload.pedal_freno}, \
        @pedal_freno_observacion = N'{payload.pedal_freno_observacion}' , \
        @pedal_freno_foto = N'{payload.pedal_freno_foto}' , \
        @pedal_estacionamiento = {payload.pedal_estacionamiento}, \
        @pedal_estacionamiento_observacion = N'{payload.pedal_estacionamiento_observacion}' , \
        @pedal_estacionamiento_foto = N'{payload.pedal_estacionamiento_foto}' , \
        @cinturon_seguridad = {payload.cinturon_seguridad}, \
        @cinturon_seguridad_observacion = N'{payload.cinturon_seguridad_observacion}' , \
        @cinturon_seguridad_foto = N'{payload.cinturon_seguridad_foto}' , \
        @cuadro_instrumentos = {payload.cuadro_instrumentos}, \
        @cuadro_instrumentos_observacion = N'{payload.cuadro_instrumentos_observacion}' ,\
        @cuadro_instrumentos_foto = N'{payload.cuadro_instrumentos_foto}' ,\
        @aire_acondicionado = {payload.aire_acondicionado},\
        @aire_acondicionado_observacion = N'{payload.aire_acondicionado_observacion}' ,\
        @aire_acondicionado_foto = N'{payload.aire_acondicionado_foto}' ,\
        @bocina_claxon = {payload.bocina_claxon},\
        @bocina_claxon_observacion = N'{payload.bocina_claxon_observacion}' ,\
        @bocina_claxon_foto = N'{payload.bocina_claxon_foto}' ,\
        @iluminacion_interior = {payload.iluminacion_interior},\
        @iluminacion_interior_observacion = N'{payload.iluminacion_interior_observacion}' ,\
        @iluminacion_interior_foto = N'{payload.iluminacion_interior_foto}' ,\
        @iluminacion_externa = {payload.iluminacion_externa},\
        @iluminacion_externa_observacion = N'{payload.iluminacion_externa_observacion}' ,\
        @iluminacion_externa_foto = N'{payload.iluminacion_externa_foto}' ,\
        @limpiaparabrisas = {payload.limpiaparabrisas}, \
        @limpiaparabrisas_observacion = N'{payload.limpiaparabrisas_observacion}' ,\
        @limpiaparabrisas_foto = N'{payload.limpiaparabrisas_foto}' , \
        @limpia_medallon = {payload.limpia_medallon}, \
        @limpia_medallon_observacion = N'{payload.limpia_medallon_observacion}' , \
        @limpia_medallon_foto = N'{payload.limpia_medallon_foto}' , \
        @neumaticos_friccion = {payload.neumaticos_friccion}, \
        @neumaticos_friccion_observacion = N'{payload.neumaticos_friccion_observacion}' ,  \
        @neumaticos_friccion_foto = N'{payload.neumaticos_friccion_foto}' , \
        @otros_vehiculo_en_piso = {payload.otros_vehiculo_en_piso}, \
        @otros_vehiculo_en_piso_observacion = N'{payload.otros_vehiculo_en_piso_observacion}' , \
        @otros_vehiculo_en_piso_foto = N'{payload.otros_vehiculo_en_piso_foto}' , \
        @estado_fugas_aceite = {payload.estado_fugas_aceite}, \
        @estado_fugas_aceite_observacion = N'{payload.estado_fugas_aceite_observacion}' , \
        @estado_fugas_aceite_foto = N'{payload.estado_fugas_aceite_foto}' , \
        @estado_nivel_calidad_lubricante_transmision = {payload.estado_nivel_calidad_lubricante_transmision}, \
        @estado_nivel_calidad_lubricante_transmision_observacion = N'{payload.estado_nivel_calidad_lubricante_transmision_observacion}' , \
        @estado_nivel_calidad_lubricante_transmision_foto = N'{payload.estado_nivel_calidad_lubricante_transmision_foto}' , \
        @estado_nivel_calidad_lubricante_diferencial = {payload.estado_nivel_calidad_lubricante_diferencial}, \
        @estado_nivel_calidad_lubricante_diferencial_observacion = N'{payload.estado_nivel_calidad_lubricante_diferencial_observacion}' , \
        @estado_nivel_calidad_lubricante_diferencial_foto = N'{payload.estado_nivel_calidad_lubricante_diferencial_foto}' , \
        @cubrepolvos_flechas = {payload.cubrepolvos_flechas}, \
        @cubrepolvos_flechas_observacion = N'{payload.cubrepolvos_flechas_observacion}' , \
        @cubrepolvos_flechas_foto = N'{payload.cubrepolvos_flechas_foto}' , \
        @componentes_direccion = {payload.componentes_direccion}, \
        @componentes_direccion_observacion = N'{payload.componentes_direccion_observacion}' , \
        @componentes_direccion_foto = N'{payload.componentes_direccion_foto}' , \
        @componentes_suspesion = {payload.componentes_suspesion}, \
        @componentes_suspesion_observacion = N'{payload.componentes_suspesion_observacion}' , \
        @componentes_suspesion_foto = N'{payload.componentes_suspesion_foto}' , \
        @sistema_escape_completo = {payload.sistema_escape_completo}, \
        @sistema_escape_completo_observacion = N'{payload.sistema_escape_completo_observacion}' , \
        @sistema_escape_completo_foto = N'{payload.sistema_escape_completo_foto}' , \
        @sistema_alimentacion_combustible = {payload.sistema_alimentacion_combustible}, \
        @sistema_alimentacion_combustible_observacion = N'{payload.sistema_alimentacion_combustible_observacion}' , \
        @sistema_alimentacion_combustible_foto = N'{payload.sistema_alimentacion_combustible_foto}' , \
        @filtro_combustible = {payload.filtro_combustible}, \
        @filtro_combustible_observacion = N'{payload.filtro_combustible_observacion}' , \
        @filtro_combustible_foto = N'{payload.filtro_combustible_foto}' , \
        @control_fugas_direccion_hidraulica = {payload.control_fugas_direccion_hidraulica}, \
        @control_fugas_direccion_hidraulica_observacion = N'{payload.control_fugas_direccion_hidraulica_observacion}' , \
        @control_fugas_direccion_hidraulica_foto = N'{payload.control_fugas_direccion_hidraulica_foto}' , \
        @otros_altura_total = {payload.otros_altura_total}, \
        @otros_altura_total_observacion = N'{payload.otros_altura_total_observacion}' , \
        @otros_altura_total_foto = N'{payload.otros_altura_total_foto}' , \
        @rodamiento_mazas_rueda = {payload.rodamiento_mazas_rueda}, \
        @rodamiento_mazas_rueda_observacion = N'{payload.rodamiento_mazas_rueda_observacion}' , \
        @rodamiento_mazas_rueda_foto = N'{payload.rodamiento_mazas_rueda_foto}' , \
        @holgura_partes_suspension_rueda = {payload.holgura_partes_suspension_rueda}, \
        @holgura_partes_suspension_rueda_observacion = N'{payload.holgura_partes_suspension_rueda_observacion}' , \
        @holgura_partes_suspension_rueda_foto = N'{payload.holgura_partes_suspension_rueda_foto}' , \
        @control_neumaticos_desgaste_presion = {payload.control_neumaticos_desgaste_presion}, \
        @control_neumaticos_desgaste_presion_observacion = N'{payload.control_neumaticos_desgaste_presion_observacion}' , \
        @control_neumaticos_desgaste_presion_foto = N'{payload.control_neumaticos_desgaste_presion_foto}' , \
        @profundidad = {payload.profundidad}, \
        @profundidad_observacion = N'{payload.profundidad_observacion}' , \
        @profundidad_foto = N'{payload.profundidad_foto}' , \
        @presion = {payload.presion}, \
        @presion_observacion = N'{payload.presion_observacion}' , \
        @presion_foto = N'{payload.presion_foto}' , \
        @otros_altura_media = {payload.otros_altura_media}, \
        @otros_altura_media_observacion = N'{payload.otros_altura_media_observacion}' , \
        @otros_altura_media_foto = N'{payload.otros_altura_media_foto}' , \
        @nivel_calidad_aceite_motor = {payload.nivel_calidad_aceite_motor}, \
        @nivel_calidad_aceite_motor_observacion = N'{payload.nivel_calidad_aceite_motor_observacion}' , \
        @nivel_calidad_aceite_motor_foto = N'{payload.nivel_calidad_aceite_motor_foto}' , \
        @filtro_aire = {payload.filtro_aire}, \
        @filtro_aire_observacion = N'{payload.filtro_aire_observacion}' , \
        @filtro_aire_foto = N'{payload.filtro_aire_foto}' , \
        @filtro_polen = {payload.filtro_polen}, \
        @filtro_polen_observacion = N'{payload.filtro_polen_observacion}' , \
        @filtro_polen_foto = N'{payload.filtro_polen_foto}' , \
        @filtro_pcv = {payload.filtro_pcv}, \
        @filtro_pcv_observacion = N'{payload.filtro_pcv_observacion}' , \
        @filtro_pcv_foto = N'{payload.filtro_pcv_foto}' , \
        @valvula_pcv = {payload.valvula_pcv}, \
        @valvula_pcv_observacion = N'{payload.valvula_pcv_observacion}' , \
        @valvula_pcv_foto = N'{payload.valvula_pcv_foto}' , \
        @bujias_encendido = {payload.bujias_encendido}, \
        @bujias_encendido_observacion = N'{payload.bujias_encendido_observacion}' , \
        @bujias_encendido_foto = N'{payload.bujias_encendido_foto}' , \
        @cables_bujias_bobinas_ignicion = {payload.cables_bujias_bobinas_ignicion}, \
        @cables_bujias_bobinas_ignicion_observacion = N'{payload.cables_bujias_bobinas_ignicion_observacion}' , \
        @cables_bujias_bobinas_ignicion_foto = N'{payload.cables_bujias_bobinas_ignicion_foto}' , \
        @nivel_anticongenlante = {payload.nivel_anticongenlante}, \
        @nivel_anticongenlante_observacion = N'{payload.nivel_anticongenlante_observacion}' , \
        @nivel_anticongenlante_foto = N'{payload.nivel_anticongenlante_foto}' , \
        @tapon_radiador = {payload.tapon_radiador}, \
        @tapon_radiador_observacion = N'{payload.tapon_radiador_observacion}' , \
        @tapon_radiador_foto = N'{payload.tapon_radiador_foto}' , \
        @mangueras_sistema = {payload.mangueras_sistema}, \
        @mangueras_sistema_observacion = N'{payload.mangueras_sistema_observacion}' , \
        @mangueras_sistema_foto = N'{payload.mangueras_sistema_foto}' , \
        @desempeño_ventilador = {payload.desempeño_ventilador}, \
        @desempeño_ventilador_observacion = N'{payload.desempeño_ventilador_observacion}' , \
        @desempeño_ventilador_foto = N'{payload.desempeño_ventilador_foto}' , \
        @calidad_liquido_limpiaparabrisas = {payload.calidad_liquido_limpiaparabrisas}, \
        @calidad_liquido_limpiaparabrisas_observacion = N'{payload.calidad_liquido_limpiaparabrisas_observacion}' , \
        @calidad_liquido_limpiaparabrisas_foto = N'{payload.calidad_liquido_limpiaparabrisas_foto}' , \
        @calidad_aceite_direccion_hidraulica = {payload.calidad_aceite_direccion_hidraulica}, \
        @calidad_aceite_direccion_hidraulica_observacion = N'{payload.calidad_aceite_direccion_hidraulica_observacion}' , \
        @calidad_aceite_direccion_hidraulica_foto = N'{payload.calidad_aceite_direccion_hidraulica_foto}' , \
        @calidad_aceite_transmision_bayoneta = {payload.calidad_aceite_transmision_bayoneta}, \
        @calidad_aceite_transmision_bayoneta_observacion = N'{payload.calidad_aceite_transmision_bayoneta_observacion}' , \
        @calidad_aceite_transmision_bayoneta_foto = N'{payload.calidad_aceite_transmision_bayoneta_foto}' , \
        @liquido_bateria_condiciones = {payload.liquido_bateria_condiciones}, \
        @liquido_bateria_condiciones_observacion = N'{payload.liquido_bateria_condiciones_observacion}' , \
        @liquido_bateria_condiciones_foto = N'{payload.liquido_bateria_condiciones_foto}' ,    \
        @bandas_poly_v = {payload.bandas_poly_v}, \
        @bandas_poly_v_observacion = N'{payload.bandas_poly_v_observacion}' , \
        @bandas_poly_v_foto = N'{payload.bandas_poly_v_foto}' , \
        @poleas_banda = {payload.poleas_banda}, \
        @poleas_banda_observacion = N'{payload.poleas_banda_observacion}' , \
        @poleas_banda_foto = N'{payload.poleas_banda_foto}' , \
        @banda_tiempo = {payload.banda_tiempo}, \
        @banda_tiempo_observacion = N'{payload.banda_tiempo_observacion}' , \
        @banda_tiempo_foto = N'{payload.banda_tiempo_foto}' , \
        @otros_habitaculo_motor = {payload.otros_habitaculo_motor}, \
        @otros_habitaculo_motor_observacion = N'{payload.otros_habitaculo_motor_observacion}' , \
        @otros_habitaculo_motor_foto = N'{payload.otros_habitaculo_motor_foto}' , \
        @reset_intervalo_servicio = {payload.reset_intervalo_servicio}, \
        @reset_intervalo_servicio_observacion = N'{payload.reset_intervalo_servicio_observacion}' , \
        @reset_intervalo_servicio_foto = N'{payload.reset_intervalo_servicio_foto}' , \
        @ajuste_tornillos_neumaticos_torquimetro = {payload.ajuste_tornillos_neumaticos_torquimetro}, \
        @ajuste_tornillos_neumaticos_torquimetro_observacion = N'{payload.ajuste_tornillos_neumaticos_torquimetro_observacion}' , \
        @ajuste_tornillos_neumaticos_torquimetro_foto = N'{payload.ajuste_tornillos_neumaticos_torquimetro_foto}' , \
        @limpiar_libricar_puertas_cerraduras = {payload.limpiar_libricar_puertas_cerraduras}, \
        @limpiar_libricar_puertas_cerraduras_observacion = N'{payload.limpiar_libricar_puertas_cerraduras_observacion}' , \
        @limpiar_libricar_puertas_cerraduras_foto = N'{payload.limpiar_libricar_puertas_cerraduras_foto}' , \
        @completar_plan_mantenimiento = {payload.completar_plan_mantenimiento}, \
        @completar_plan_mantenimiento_observacion = N'{payload.completar_plan_mantenimiento_observacion}' , \
        @completar_plan_mantenimiento_foto = N'{payload.completar_plan_mantenimiento_foto}' , \
        @fecha = N'{payload.Fecha}' , \
        @Id_empleado = {payload.IdEmpleado} , \
        @Id_vehiculo = {payload.IdVehiculo} , \
        @id_ordendeservicio = {payload.Id_ordendeservicio} , \
        @NumeroSerie = '{payload.NumeroSerie}' """

    try:
        with engine.begin() as conn:
            # Ejecutar la consulta SQL
            roles_df = pd.read_sql(query, conn)

        # Crear el objeto de respuesta
        response = ResponseModel(id_resultado=1, respuesta="El checklist se actualizó correctamente.")
        return JSONResponse(status_code=200, content=response.dict())

    except Exception as e:
        # Manejo de errores
        raise HTTPException(status_code=500, detail=str(e))




@app.get(
        path="/api/obtenerchecklist",
        name='Obtener checklist',
        tags=['Checklist'],
        description='Método para obtener la informacion de 1 checklist',
        response_model=Checklist
)
def getempleados(Idchecklist: int):
    query = f"exec [dbo].[sp_get_all_checklist] @IdCheckList = {Idchecklist}"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])


@app.get(
        path="/api/obtenertecnicos",
        name='Obtener tecnicos',
        tags=['Tecnicos'],
        description='Método para obtener la informacion todos los técnicos',
        response_model=Tecnicos
)
def getchecklists():
    query = f"exec [dbo].[ObtenerTecnicos]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/obtenerallchecklists",
        name='Obtener todos los checklists',
        tags=['Checklist'],
        description='Método para obtener la informacion todos los checklists',
        response_model=Checklist
)
def getallchecklists():
    query = f"exec [dbo].[sp_get_all_checklistv2]"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)

@app.get(
        path="/api/obtenerchecklists",
        name='Obtener checklists',
        tags=['Checklist'],
        description='Método para obtener la informacion todos los checklist',
        response_model=Checklist
)
def getchecklists(IdEmpleado:int):
    query = f"exec [dbo].[ObtenerCheckLists] @IdEmpleado = {IdEmpleado} "
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)
@app.get(
    path="/api/obteneridOrden",
    name="Obtener ID de Orden de Servicio",
    tags=["Orden"],
    description="Obtiene el siguiente ID secuencial para la orden de servicio."
)
def obtener_id_orden():
    query = f"exec [dbo].ObtenerUltimoIdOrdenServicio"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    print(resultado)
    return JSONResponse(status_code=200,content=resultado[0])

@app.get(
    path="/api/obteneridCheck",
    name='Obtener IDs del Checklist',
    tags=['Checklist'],
    description='Obtiene todos los IDs de los checklists existentes',
)
def obtener_id_check():
    query = "exec [dbo].ObtenerIdCheckList"
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200, content=resultado)


@app.get(
    path="/api/obtenerchecklisthtml",
    name='Obtener checklist HTML',
    tags=['Checklist'],
    description='Método para obtener el HTML del checklist',
)
def obtener_checklist_html(Idchecklist: int):

    query = f"exec [dbo].[sp_get_all_checklist] @IdCheckList = {Idchecklist}"
    with engine.begin() as conn:
      conn.execution_options(autocommit=True)
    roles_df = pd.read_sql(query, conn)

        # Asegúrate de que hay al menos dos filas en el DataFrame

            # Accede al valor en la segunda fila de la columna 'Nombre'


            # Construye el HTML con el valor obtenido
    id = roles_df['id'].iloc[0]
    fecha = ['fecha'].iloc[0]
    idEmpleado = ['Id_empleado'].iloc[0]
    idVehiculo = ['Id_vehiculo'].iloc[0]
    idOrden= ['id_ordendeservicio'].iloc[0]
    lecturaCodigo = ['lectura_codigos'].iloc[0]
    lecturaCodigoObservacion =['lectura_codigos_observacion'].iloc[0]
    servofreno = ['servofreno'].iloc[0]
    servofrenoObservacion =['servofreno_observacion'].iloc[0]
    pedalFreno =['pedal_freno'].iloc[0]
    pedalFrenoObservacion =['pedal_freno_observacion'].iloc[0]
    pedalEstacionamiento =['pedal_estacionamiento'].iloc[0]
    pedalEstacionamientoObservacion=['pedal_estacionamiento_observacion'].iloc[0]
    cinturonSeguridad =['cinturon_seguridad'].iloc[0]
    cinturonSeguridadObservacion=['cinturon_seguridad_observacion'].iloc[0]
    cuadroInstrumentos=['cuadro_instrumentos'].iloc[0]
    cuadroInstrumentosObservacion=['cuadro_instrumentos_observacion'].iloc[0]
    aireAcondicionado =['aire_acondicionado'].iloc[0]
    aireAcondicionadoObservacion=['aire_acondicionado_observacion'].iloc[0]
    bocinaClaxon=['bocina_claxon'].iloc[0]
    bocinaClaxonObservacion=['bocina_claxon_observacion'].iloc[0]
    iluminacionInterior=['iluminacion_interior'].iloc[0]
    iluminacionInteriorObservacion=['iluminacion_interior_observacion'].iloc[0]
    iluminacionExterna=['iluminacion_externa'].iloc[0]
    iluminacionExternaObservacion=['iluminacion_externa_observacion'].iloc[0]
    limpiaparabrisas=['limpiaparabrisas'].iloc[0]
    limpiaparabrisasObservacion=['limpiaparabrisas_observacion'].iloc[0]
    limpiaMedallon=['limpia_medallon'].iloc[0]
    limpiaMedallonObservacion=['limpia_medallon_observacion'].iloc[0]
    neumaticosFriccion=['neumaticos_friccion'].iloc[0]
    neumaticosFriccionObservacion=['neumaticos_friccion_observacion'].iloc[0]
    otroVehiculosPiso=['otros_vehiculo_en_piso'].iloc[0]
    otroVehiculosPisoObservacion=['otros_vehiculo_en_piso_observacion'].iloc[0]
    estadoFugasAceite=['estado_fugas_aceite'].iloc[0]
    estadoFugasAceiteObservacion=['estado_fugas_aceite_observacion'].iloc[0]
    estadoNivelCalidad=['estado_nivel_calidad_lubricante_transmision'].iloc[0]
    estadoNivelCalidadObservacion=['estado_nivel_calidad_lubricante_transmision_observacion'].iloc[0]
    estadoNivelCalidadDiferencial=['estado_nivel_calidad_lubricante_diferencial'].iloc[0]
    estadoNivelCalidadDiferencialObservacion=['estado_nivel_calidad_lubricante_diferencial_observacion'].iloc[0]
    cubrepolvosFlechas = ['cubrepolvos_flechas'].iloc[0]
    cubrepolvosFlechasObservacion=['cubrepolvos_flechas_observacion'].iloc[0]
    componentesDireccion=['componentes_direccion'].iloc[0]
    componentesDireccionObservacion=['componentes_direccion_observacion'].iloc[0]
    componentesSuspension=['componentes_suspesion'].iloc[0]
    componentesSuspensionObservacion=['componentes_suspesion_observacion'].iloc[0]
    sistemasEscape=['sistema_escape_completo'].iloc[0]
    sistemasEscapeObservacion=['sistema_escape_completo_observacion'].iloc[0]
    sistemaAlimentacion=['sistema_alimentacion_combustible'].iloc[0]
    sistemaAlimentacionObservacion=['sistema_alimentacion_combustible_observacion'].iloc[0]
    filtroCombustible=['filtro_combustible'].iloc[0]
    filtroCombustibleObservacion=['filtro_combustible_observacion'].iloc[0]
    controlFugasDireccion=['control_fugas_direccion_hidraulica'].iloc[0]
    controlFugasDireccionObservacion=['control_fugas_direccion_hidraulica_observacion'].iloc[0]
    otroAltura=['otros_altura_total'].iloc[0]
    otrosAlturaObservacion=['otros_altura_total_observacion'].iloc[0]
    rodamientoMazas=['rodamiento_mazas_rueda'].iloc[0]
    rodamientoMAzasObservacion=['rodamiento_mazas_rueda_observacion'].iloc[0]
    holguraSuspension=['holgura_partes_suspension_rueda'].iloc[0]
    holguraSuspensionObservacion=['holgura_partes_suspension_rueda_observacion'].iloc[0]
    controlNeumaticos=['control_neumaticos_desgaste_presion'].iloc[0]
    controlNeumaticosObservacion=['control_neumaticos_desgaste_presion_observacion'].iloc[0]
    profundidad=['profundidad'].iloc[0]
    profundidadObservacion=['profundidad_observacion'].iloc[0]
    presion=['presion'].iloc[0]
    presionObservacion=['presion_observacion'].iloc[0]
    otrosAltura=['otros_altura_media'].iloc[0]
    otrosAlturaObservacion=['otros_altura_media_observacion'].iloc[0]
    nivelCalidadAceite=['nivel_calidad_aceite_motor'].iloc[0]
    nivelCalidadAceiteObservacion=['nivel_calidad_aceite_motor_observacion'].iloc[0]
    filtroAire=['filtro_aire'].iloc[0]
    filtroAireObservacion=['filtro_aire_observacion'].iloc[0]
    filtroPolen=['filtro_polen'].iloc[0]
    filtroPolenObservacion=['filtro_polen_observacion'].iloc[0]
    filtroPcv=['filtro_pcv'].iloc[0]
    filtroPcvObservacion=['filtro_pcv_observacion'].iloc[0]
    valvulaPcv=['valvula_pcv'].iloc[0]
    valvulaPcvObservacion=['valvula_pcv_observacion'].iloc[0]
    bujiasEncendido=['bujias_encendido'].iloc[0]
    bujiasEncendidoObservacion=['bujias_encendido_observacion'].iloc[0]
    cablesBujiasBobinas=['cables_bujias_bobinas_ignicion'].iloc[0]
    cablesBujiasBobinasObservacion=['cables_bujias_bobinas_ignicion_observacion'].iloc[0]
    nivelAnticongelante=['nivel_anticongenlante'].iloc[0]
    nivelAnticongelanteObservacion=['nivel_anticongenlante_observacion'].iloc[0]
    taponRadiador=['tapon_radiador'].iloc[0]
    taponRadiadorObservacion=['tapon_radiador_observacion'].iloc[0]
    manguerasSistema=['mangueras_sistema'].iloc[0]
    manguerasSistemaObservacion=['mangueras_sistema_observacion'].iloc[0]
    desempeñoVentilador=['desempeño_ventilador'].iloc[0]
    desempeñoVentiladorObservacion=['desempeño_ventilador_observacion'].iloc[0]
    calidadLiquidoLimpia=['calidad_liquido_limpiaparabrisas'].iloc[0]
    calidadLiquidoLimpiaObservacion=['calidad_liquido_limpiaparabrisas_observacion'].iloc[0]
    calidadAceiteDireccion=['calidad_aceite_direccion_hidraulica'].iloc[0]
    calidadAceiteDireccionObservacion=['calidad_aceite_direccion_hidraulica_observacion'].iloc[0]
    calidadAceiteTransmision=['calidad_aceite_transmision_bayoneta'].iloc[0]
    calidadAceiteTransmisionObservacion=['calidad_aceite_transmision_bayoneta_observacion'].iloc[0]
    liquidoBateriaCondiciones=['liquido_bateria_condiciones'].iloc[0]
    liquidoBateriaCondicionesObservacion=['liquido_bateria_condiciones_observacion'].iloc[0]
    bandasPoly=['bandas_poly_v'].iloc[0]
    bandasPolyObservacion=['bandas_poly_v_observacion'].iloc[0]
    poleasBanda=['poleas_banda'].iloc[0]
    poleasBandaObservacion=['poleas_banda_observacion'].iloc[0]
    bandaTiempo=['banda_tiempo'].iloc[0]
    bandaTiempoObservacion=['banda_tiempo_observacion'].iloc[0]
    otrosHabitaculo=['otros_habitaculo_motor'].iloc[0]
    otrosHabitaculoObservacion=['otros_habitaculo_motor_observacion'].iloc[0]
    resetIntervalo=['reset_intervalo_servicio'].iloc[0]
    resetIntervaloObservacion=['reset_intervalo_servicio_observacion'].iloc[0]
    ajusteTornillosNeumaticos=['ajuste_tornillos_neumaticos_torquimetro'].iloc[0]
    ajusteTornillosNeumaticosObservacion=['ajuste_tornillos_neumaticos_torquimetro_observacion'].iloc[0]
    limpiarLubricarPuertas=['limpiar_libricar_puertas_cerraduras'].iloc[0]
    limpiarLubricarPuertasObservacion=['limpiar_libricar_puertas_cerraduras_observacion'].iloc[0]
    completarPlanMantenimiento=['completar_plan_mantenimiento'].iloc[0]
    completarPlanMantenimientoObservacion=['completar_plan_mantenimiento_observacion'].iloc[0]
    NumeroSerie=['NumeroSerie'].iloc[0]

    htmlstring = r"""html:"<!DOCTYPE html> <html lang="es"> <head>
      <meta charset="UTF-8">
      <meta name="viewport" content="width=device-width, initial-scale=1.0">
      <title>Check List General</title>
      <style>
        body {
            font-family: Arial, sans-serif;
        }
        h1, h2 {
            text-align: center;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 20px;
        }
        td, th {
            padding: 8px;
            text-align: left;
            border: 1px solid #ddd;
        }
        th {
            background-color: #f2f2f2;
        }
        .highlight {
            background-color: #f9e44c;
        }
        .footer {
            text-align: center;
            margin-top: 20px;
        }
     </style>
      </head>
    <body>"""
    htmlstring=f"""
      <h1>CHECK LIST GENERAL</h1>
      <h2>Aplicable para todos los servicios</h2>

      <table>
        <tr>
            <th>No. de Orden</th>
            <td>{idOrden}</td>
            <th>No. Folio</th>
            <td>484</td>
        </tr>
        <tr>
            <th>No. Serie del Vehículo</th>
            <td colspan="3">{NumeroSerie}</td>
        </tr>
        <tr>
            <th>Fecha</th>
            <td colspan="3">{fecha}</td>
        </tr>
      </table>

      <h3>Checklist</h3>

      <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr class="highlight">
                <td>Vehículo en el piso</td>
                <td>OK</td>
                <td>NO OK</td>
                <td></td>
            </tr>
            <tr>
                <td>Lectura de códigos de falla. (Nº de averías en sus sistemas)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{lecturaCodigoObservacion}</td>
            </tr>
            <tr>
                <td>Funcionamiento del servofreno</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{servofrenoObservacion}</td>
            </tr>
            <tr>
                <td>Funcionamiento del pedal de freno</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{pedalFrenoObservacion}</td>
            </tr>
            <tr>
                <td>Funcionamiento del pedal de estacionamiento</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{pedalEstacionamientoObservacion}</td>
            </tr>
            <tr>
                <td>Cuadro de instrumentos (funcionamiento, iluminación, testigos)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{cuadroInstrumentosObservacion}</td>
            </tr>
            <tr>
                <td>Iluminación interior completa (cortesía, vanidad, guantera, etc.)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{iluminacionInteriorObservacion}</td>
            </tr>
            <tr>
                <td>Iluminación exterior (cortes, bajas, altas, reversa, niebla, nivel)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{iluminacionExternaObservacion}</td>
            </tr>
            <tr>
                <td>Sistema limpiaparabrisas completo, chorritos, plumas</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{limpiaparabrisasObservacion}</td>
            </tr>
            <tr>
                <td>Sistema limpia medallón completo, chorritos, plumas</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{limpiaMedallonObservacion}</td>
            </tr>
            <tr>
                <td>Sistema de frenos (desgaste y presión)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{neumaticosFriccionObservacion}</td>
            </tr>
        </tbody>
    </table>

    <h3>Vehículo en altura total</h3>

    <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Estado de fugas de aceite de motor y transmisión</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{estadoFugasAceiteObservacion}</td>
            </tr>
            <tr>
                <td>Estado de nivel y calidad de lubricantes en transmisión</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{estadoNivelCalidadObservacion}</td>
            </tr>
            <tr>
                <td>Estado de nivel y calidad de lubricantes en diferencial</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{estadoNivelCalidadDiferencialObservacion}</td>
            </tr>
            <tr>
                <td>Estado de cubrepolvos de flechas homocinéticas</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{cubrepolvosFlechasObservacion}</td>
            </tr>
            <tr>
                <td>Estado de componentes de dirección</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{componentesDireccionObservacion}</td>
            </tr>
            <tr>
                <td>Estado de suspensión</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{componentesSuspensionObservacion}</td>
            </tr>
            <tr>
                <td>Estado del sistema de alimentación de combustible (fugas)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{sistemaAlimentacionObservacion}</td>
            </tr>
            <tr>
                <td>Control de fugas de dirección hidráulica</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{controlFugasDireccionObservacion}</td>
            </tr>
        </tbody>
    </table>

    <h3>Vehículo en altura media</h3>

    <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Control de rodamiento y mazas de rueda</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{rodamientoMAzasObservacion}</td>
            </tr>
            <tr>
                <td>Holgura/juego en partes de suspensión y dirección</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{holguraSuspensionObservacion}</td>
            </tr>
            <tr>
                <td>Profundidad (D.D)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{profundidadObservacion}</td>
            </tr>
            <tr>
                <td>Presión (D.D)</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{presionObservacion}</td>
            </tr>
        </tbody>
    </table>

    <h3>Control de Habitáculo de motor</h3>

    <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Nivel y calidad de aceite de motor</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{nivelCalidadAceiteObservacion}</td>
            </tr>
            <tr>
                <td>Filtro de aire</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{filtroAireObservacion}</td>
            </tr>
        </tbody>
    </table>

    <h3>Revisión Final</h3>

    <table>
        <thead>
            <tr>
                <th>Item</th>
                <th>OK</th>
                <th>NO OK</th>
                <th>Observaciones</th>
            </tr>
        </thead>
        <tbody>
            <tr>
                <td>Reset de intervalo de servicio</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{resetIntervaloObservacion}</td>
            </tr>
            <tr>
                <td>Ajuste y verificación de alineación</td>
                <td>OK</td>
                <td>NO OK</td>
                <td>{ajusteTornillosNeumaticosObservacion}</td>
            </tr>
            <tr>
                <td>Ajuste de frenos</td>
                <td>OK</td>
                <td>NO OK</td>
                <td></td>
            </tr>
            <tr>
                <td>Revisión de luces</td>
                <td>OK</td>
                <td>NO OK</td>
                <td></td>
            </tr>
        </tbody>
    </table>

    <div class="footer">
        <p>Firma Cliente: ___________________________</p>
        <p>Firma Responsable: ___________________________</p>
    </div>
    </body>
    </html>"""
    img = "\\img1.jpg"
    pdf_path = "example.pdf"
          # Rutas y configuraciones para Linux
    path_wkhtmltopdf = '/usr/local/bin/wkhtmltopdf'
    config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)
    pdfkit.from_string(htmlstring, 'reporte.pdf', configuration=config)
    return JSONResponse(content={"message": "PDF creado exitosamente"}, status_code=200)
    return JSONResponse(content={"error": str(e)}, status_code=500)

@app.get(
        path="/api/historicocheck",
        name='Obtener clientes',
        tags=['Historico'],
        description='Método para obtener la informacion de todos los clientes',
        response_model=List[GetCliente]
)
def getclientes(busqueda = ""):
    query = f"exec Clientes.clienteinsupdel @Accion = 2,@ParametroBusqueda = '{busqueda}' "
    roles_df = pd.read_sql(query, engine)
    resultado = roles_df.to_dict(orient="records")
    return JSONResponse(status_code=200,content=resultado)
@app.get(
        path="/api/clienteexiste",
        name='Obtener existencia de cliente',
        tags=['Clientes'],
        description='Método para obtener la informacion de todos los clientes',
        response_model=List[GetCliente]
)
def cliente_existe(nombre: str = "", email: str = ""):
    conn = engine.raw_connection()
    try:
        cursor = conn.cursor()
        cursor.execute("""
            DECLARE @Existe BIT;
            EXEC sp_Cliente_Existe @Nombre = ?, @Email = ?, @Existe = @Existe OUTPUT;
            SELECT @Existe AS Existe;
        """, nombre, email)

        row = cursor.fetchone()
        existe_valor = row[0] if row else 0

        return JSONResponse(status_code=200, content={"existe": bool(existe_valor)})
    finally:
        conn.close()


@app.post(
        path="/api/ordenservice",
        name='Insertar orden de servicio',
        tags=['Orden'],
        description='Método para insertar la orden de servicio',
        response_model=OrdenService
)
def saveordenservice(payload: OrdenService):
    query = f"""EXEC InsertarOrdenServicio @idCliente = {payload.IdCliente} , \
        @idEmpleado = '{payload.IdEmpleado}' """
    print(query)

    with engine.begin() as conn:
        conn.execution_options(autocommit = True)
        roles_df = pd.read_sql(query, conn)
    dumpp = ResponseModel(id_resultado=1,respuesta="La orden se ha guardado de manera correcta")
    dict = dumpp.model_dump()
    return JSONResponse(status_code=200, content=dict)


@app.post(
    path="/api/asignarTecnicoAOrden",
    name="Asignar orden de servicio",
    tags=["AsignarOrden"],
    description="Método para asignar un técnico a una orden de servicio",
    response_model=ResponseModel,
)
def saveAsignacion(payload: AsignarOrden):
    # Creación de la consulta SQL con parámetros de manera segura
    query = text("""
        EXEC AsignarTecnicoAOrden :IdOrden, :IdTecnico
    """)

    # Conexión y ejecución del query
    with engine.begin() as conn:
        conn.execution_options(autocommit=True)

        # Ejecutando la consulta con los parámetros adecuados
        result = conn.execute(query, {"IdOrden": payload.IdOrden, "IdTecnico": payload.IdTecnico})

    # Aquí deberías verificar si la ejecución fue exitosa o si hubo algún error
    # Para este caso, vamos a asumir que la respuesta del SP incluye un mensaje
    # En este ejemplo se supone que la respuesta será capturada por el procedimiento
    response_data = {
        "id_resultado": 1,
        "respuesta": "La orden se ha guardado de manera correcta"
    }

    return JSONResponse(status_code=200, content=response_data)
@app.get(
        path="/api/obtenerreporteporId",
        name='Obtener reporte de venta por id',
        tags=['ReporteVentas'],
        description='Método para obtener la informacion de 1 reporte de ventas',
        response_model=ReporteVentas
)
def getreporteporId(IdReporte: int):
    try:
        query = f"exec [dbo].[ObtenerReporteVentasPorID] @IdReporte = {IdReporte}"
        roles_df = pd.read_sql(query, engine)

        # Convertir todo el DataFrame a cadenas antes de convertir a dict
        resultado = roles_df.astype(str).to_dict(orient="records")

        return JSONResponse(status_code=200, content=resultado[0])
    except Exception as e:
        # Respuesta de error
        raise HTTPException(
            status_code=500,
            detail=f"Error al obtener el reporte: {str(e)}",
        )

@app.get(
    path="/api/obtenerreportes",
    name='Obtener todos los reportes de venta',
    tags=['ReporteVentas'],
    description='Método para obtener la información de todos los reportes de venta',
    response_model=List[ReporteVentas]  # Ajusta según corresponda
)
def getallreportes():
    try:
        query = "exec [dbo].[ObtenerAllReporteVentas]"
        roles_df = pd.read_sql(query, engine)

        # Convertir columnas datetime a cadenas en formato ISO
        for col in roles_df.select_dtypes(include=["datetime64[ns]", "datetime64[ns, UTC]"]).columns:
            roles_df[col] = roles_df[col].dt.strftime('%Y-%m-%dT%H:%M:%S')

        # Convertir a lista de diccionarios
        resultado = roles_df.to_dict(orient="records")

        # Serializar el contenido con jsonable_encoder para manejar tipos complejos
        return JSONResponse(status_code=200, content=jsonable_encoder(resultado))
    except Exception as e:
        # Respuesta de error
        raise HTTPException(
            status_code=500,
            detail=f"Error al obtener los reportes: {str(e)}",
        )

@app.post(
    path="/api/reporteventas",
    name='Insertar reporte',
    tags=['ReporteVentas'],
    description='Método para insertar el reporte de ventas',
    response_model=ReporteVentas
)
def savereporteventas(payload: ReporteVentas):
    # Query SQL con placeholders adecuados para SQL Server
    try:
        # Crear el diccionario de parámetros desde el payload
        parametros = payload.dict()
        query = text("""
        EXEC InsertServiceForm
        @date = :date,
        @service_order_id = :service_order_id,
        @vehicle_id = :vehicle_id,
        @credit = :credit,
        @initial_service = :initial_service,
        @finalized = :finalized,
        @reception = :reception,
        @entry = :entry,
        @repair = :repair,
        @checklist = :checklist,
        @technician = :technician,
        @quotation = :quotation,
        @authorization = :authorization,
        @additional = :additional,
        @washing = :washing,
        @delivery = :delivery,
        @comments = :comments
    """)

     # Ejecutar la consulta SQL
        with engine.begin() as conn:
            conn.execute(query, parametros)

        # Respuesta de éxito
        return JSONResponse(
            status_code=200,
            content={
                "id_resultado": 1,
                "respuesta": "El reporte de venta se guardó de manera correcta",
                "detalles": parametros,
            },
        )

    except Exception as e:
        # Respuesta de error
        raise HTTPException(
            status_code=500,
            detail=f"Error al guardar el reporte: {str(e)}",
        )
if __name__ == '__main__':
    app.run()